"""
Extraction automatique du Reglement de Consultation (RC) depuis les fichiers DCE.

Identifie le RC, extrait le texte, puis parse les informations cles :
- Pieces exigees
- Criteres d'attribution avec ponderations
- Date limite de remise
- Modalites de remise
- Variantes autorisees
- Lots
- Conditions de participation
- Duree du marche
- Montant estime
"""

import re
import logging
from pathlib import Path
from datetime import datetime

logger = logging.getLogger("ao_hunter.extraction_rc")

# Patterns pour identifier le fichier RC dans le DCE
PATTERNS_RC = [
    re.compile(r"(?:^|[\s_\-])RC(?:[\s_\-\.]|$)", re.IGNORECASE),
    re.compile(r"reglement.*consultation", re.IGNORECASE),
    re.compile(r"r[eè]glement.*consultation", re.IGNORECASE),
    re.compile(r"RC[\s_\-]?(?:AO|marche|consultation)", re.IGNORECASE),
    re.compile(r"reglement.*particulier", re.IGNORECASE),
]

# Extensions lisibles
EXTENSIONS_LISIBLES = {".pdf", ".docx"}


def _identifier_rc(dossier_dce: Path) -> Path | None:
    """Identifie le fichier RC dans un dossier DCE.

    Returns:
        Path du fichier RC ou None si non trouve.
    """
    if not dossier_dce.exists() or not dossier_dce.is_dir():
        return None

    fichiers = sorted(
        f for f in dossier_dce.iterdir()
        if f.is_file() and f.suffix.lower() in EXTENSIONS_LISIBLES
    )

    # Passe 1 : matching exact sur le nom
    for fichier in fichiers:
        nom = fichier.stem
        for pattern in PATTERNS_RC:
            if pattern.search(nom):
                logger.info(f"RC identifie: {fichier.name}")
                return fichier

    # Passe 2 : si un seul PDF, le prendre comme RC probable
    pdfs = [f for f in fichiers if f.suffix.lower() == ".pdf"]
    if len(pdfs) == 1:
        logger.info(f"Un seul PDF dans le DCE, utilise comme RC probable: {pdfs[0].name}")
        return pdfs[0]

    logger.info(f"RC non identifie dans {dossier_dce} ({len(fichiers)} fichiers)")
    return None


def _extraire_texte(fichier: Path) -> str:
    """Extrait le texte d'un fichier PDF ou DOCX."""
    if fichier.suffix.lower() == ".pdf":
        return _extraire_pdf(fichier)
    elif fichier.suffix.lower() == ".docx":
        return _extraire_docx(fichier)
    return ""


def _extraire_pdf(fichier: Path) -> str:
    """Extrait le texte d'un PDF avec pypdf."""
    try:
        from pypdf import PdfReader
    except ImportError:
        logger.warning("pypdf non installe - extraction PDF impossible")
        return ""

    try:
        reader = PdfReader(str(fichier))
        pages_texte = []
        for page in reader.pages:
            texte = page.extract_text()
            if texte:
                pages_texte.append(texte)
        return "\n".join(pages_texte)
    except Exception as e:
        logger.warning(f"Erreur extraction PDF {fichier.name}: {e}")
        return ""


def _extraire_docx(fichier: Path) -> str:
    """Extrait le texte d'un DOCX avec python-docx ou zipfile."""
    try:
        from docx import Document
        doc = Document(str(fichier))
        paragraphes = [p.text for p in doc.paragraphs if p.text.strip()]
        # Aussi extraire le texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphes.append(cell.text.strip())
        return "\n".join(paragraphes)
    except ImportError:
        pass
    except Exception as e:
        logger.warning(f"Erreur extraction DOCX avec python-docx {fichier.name}: {e}")

    # Fallback : extraction basique via zipfile
    try:
        import zipfile
        with zipfile.ZipFile(str(fichier), "r") as z:
            if "word/document.xml" not in z.namelist():
                return ""
            xml_content = z.read("word/document.xml").decode("utf-8", errors="ignore")
            texte_parts = re.findall(r"<w:t[^>]*>(.*?)</w:t>", xml_content)
            return " ".join(texte_parts)
    except Exception as e:
        logger.warning(f"Erreur extraction DOCX {fichier.name}: {e}")
        return ""


# --- Parsers pour chaque information cle ---

def _extraire_pieces_exigees(texte: str) -> list[str]:
    """Extrait la liste des pieces/documents exiges."""
    pieces = []

    # Patterns pour sections "pieces a fournir"
    patterns_section = [
        re.compile(
            r"(?:pi[eè]ces?\s*[àa]\s*(?:fournir|produire|remettre|joindre)|"
            r"documents?\s*[àa]\s*(?:fournir|produire|remettre)|"
            r"composition\s*(?:du|de\s*l.)?\s*(?:dossier|offre|candidature)|"
            r"contenu\s*(?:du|de\s*l.)?\s*(?:dossier|offre))"
            r"(.*?)(?=\n\s*(?:article|\d+[\.\)]\s*[A-Z]|ARTICLE|CHAPITRE))",
            re.IGNORECASE | re.DOTALL,
        ),
    ]

    for pattern in patterns_section:
        matches = pattern.findall(texte)
        for match in matches:
            # Extraire les items de la liste
            lignes = match.split("\n")
            for ligne in lignes:
                ligne = ligne.strip()
                if not ligne or len(ligne) < 5:
                    continue
                # Lignes qui commencent par un tiret, numero, lettre)
                if re.match(r"^[\-\•\*]\s+|^\d+[\.\)]\s+|^[a-z]\)\s+", ligne):
                    piece = re.sub(r"^[\-\•\*\d\.\)a-z]+\s*", "", ligne).strip()
                    if piece and len(piece) > 3:
                        pieces.append(piece)

    # Chercher aussi des mentions specifiques de documents courants
    docs_courants = [
        (r"DC\s*1", "DC1 - Lettre de candidature"),
        (r"DC\s*2", "DC2 - Declaration du candidat"),
        (r"DC\s*4", "DC4 - Declaration de sous-traitance"),
        (r"DUME", "DUME - Document Unique de Marche Europeen"),
        (r"m[eé]moire\s*technique", "Memoire technique"),
        (r"lettre\s*de\s*candidature", "Lettre de candidature"),
        (r"(?:BPU|bordereau\s*(?:de\s*)?prix)", "Bordereau de Prix Unitaires (BPU)"),
        (r"DPGF", "DPGF - Decomposition du Prix Global et Forfaitaire"),
        (r"acte\s*d.engagement", "Acte d'engagement"),
        (r"RIB|relev[eé]\s*d.identit[eé]\s*bancaire", "RIB"),
        (r"attestation\s*(?:d.)?assurance", "Attestation d'assurance"),
        (r"Kbis|extrait\s*K", "Extrait Kbis"),
        (r"attestation\s*(?:fiscale|sociale|URSSAF)", "Attestations fiscales et sociales"),
        (r"certificat\s*(?:de\s*)?(?:capacit[eé]|qualification)", "Certificat de qualification"),
        (r"planning|calendrier\s*(?:pr[eé]visionnel)?", "Planning previsionnel"),
        (r"r[eé]f[eé]rences?\s*(?:clients?|similaires?)", "References clients"),
        (r"CV|curriculum", "CV des intervenants"),
        (r"Qualiopi", "Certificat Qualiopi"),
        (r"programme\s*(?:de\s*)?formation", "Programme de formation"),
    ]

    pieces_detectees = set(p.lower() for p in pieces)
    for pattern, nom in docs_courants:
        if re.search(pattern, texte, re.IGNORECASE):
            if nom.lower() not in pieces_detectees:
                pieces.append(nom)
                pieces_detectees.add(nom.lower())

    return pieces


def _extraire_criteres_attribution(texte: str) -> list[dict]:
    """Extrait les criteres d'attribution avec leur ponderation."""
    criteres = []

    # Pattern : "critere X : description ... XX%" ou "critere X (XX%)"
    pattern_critere = re.compile(
        r"(?:crit[eè]re\s*(?:n[o°]?\s*)?\d*\s*[\-:\.]?\s*)"
        r"(.*?)"
        r"(?:\s*[\-:]\s*|\s+)"
        r"(\d{1,3})\s*(?:%|points?|pts?)",
        re.IGNORECASE,
    )
    for match in pattern_critere.finditer(texte):
        nom = match.group(1).strip().rstrip(":-. ")
        poids = int(match.group(2))
        if nom and 0 < poids <= 100:
            criteres.append({"nom": nom, "poids": poids})

    # Pattern plus simple : "description : XX%" ou "description (XX%)"
    if not criteres:
        pattern_simple = re.compile(
            r"(?:^|\n)\s*[\-\•\*]?\s*"
            r"((?:valeur|prix|qualit[eé]|technique|m[eé]thodologie|d[eé]lai|"
            r"comp[eé]tence|exp[eé]rience|r[eé]f[eé]rence|moyens?|"
            r"p[eé]dagog|formation|innovation|environnement|social|"
            r"d[eé]veloppement\s*durable|performance)[^:\n]{0,60}?)"
            r"\s*[\-:=\(]\s*(\d{1,3})\s*(?:%|points?|pts?|\))",
            re.IGNORECASE,
        )
        for match in pattern_simple.finditer(texte):
            nom = match.group(1).strip().rstrip(":-. ")
            poids = int(match.group(2))
            if nom and 0 < poids <= 100:
                criteres.append({"nom": nom, "poids": poids})

    # Pattern pour "Prix : XX / Technique : XX"
    if not criteres:
        pattern_slash = re.compile(
            r"(prix|valeur\s*technique|qualit[eé]|technique|m[eé]thodologie|d[eé]lai)"
            r"\s*[\-:=]\s*(\d{1,3})\s*(?:%|points?|pts?|/\s*100)",
            re.IGNORECASE,
        )
        for match in pattern_slash.finditer(texte):
            nom = match.group(1).strip()
            poids = int(match.group(2))
            if 0 < poids <= 100:
                criteres.append({"nom": nom.capitalize(), "poids": poids})

    # Deduplication
    seen = set()
    unique = []
    for c in criteres:
        key = c["nom"].lower()
        if key not in seen:
            seen.add(key)
            unique.append(c)
    return unique


def _extraire_date_limite(texte: str) -> str | None:
    """Extrait la date et heure limite de remise des offres."""
    patterns = [
        # "date limite : 15/03/2026 a 12h00"
        re.compile(
            r"date\s*(?:et\s*heure\s*)?limite.*?"
            r"(\d{1,2})[/\-\.](\d{1,2})[/\-\.](\d{4})"
            r"(?:\s*[àa]\s*(\d{1,2})\s*[hH:]\s*(\d{2})?)?",
            re.IGNORECASE,
        ),
        # "avant le 15 mars 2026 a 12h00"
        re.compile(
            r"(?:avant\s*le|au\s*plus\s*tard\s*le)\s*"
            r"(\d{1,2})\s*"
            r"(janvier|f[eé]vrier|mars|avril|mai|juin|juillet|ao[uû]t|septembre|octobre|novembre|d[eé]cembre)"
            r"\s*(\d{4})"
            r"(?:\s*[àa]\s*(\d{1,2})\s*[hH:]\s*(\d{2})?)?",
            re.IGNORECASE,
        ),
        # "15/03/2026 12:00"
        re.compile(
            r"remise.*?(\d{1,2})[/\-\.](\d{1,2})[/\-\.](\d{4})"
            r"(?:\s*[àa]?\s*(\d{1,2})\s*[hH:]\s*(\d{2})?)?",
            re.IGNORECASE,
        ),
    ]

    mois_fr = {
        "janvier": "01", "fevrier": "02", "février": "02", "mars": "03",
        "avril": "04", "mai": "05", "juin": "06", "juillet": "07",
        "aout": "08", "août": "08", "septembre": "09", "octobre": "10",
        "novembre": "11", "decembre": "12", "décembre": "12",
    }

    for pattern in patterns:
        match = pattern.search(texte)
        if match:
            groups = match.groups()
            try:
                if groups[1] in mois_fr or groups[1].lower() in mois_fr:
                    # Format "15 mars 2026"
                    jour = groups[0].zfill(2)
                    mois = mois_fr.get(groups[1].lower(), groups[1])
                    annee = groups[2]
                else:
                    # Format "15/03/2026"
                    jour = groups[0].zfill(2)
                    mois = groups[1].zfill(2)
                    annee = groups[2]

                heure = groups[3].zfill(2) if groups[3] else "12"
                minute = groups[4].zfill(2) if groups[4] else "00"

                return f"{annee}-{mois}-{jour}T{heure}:{minute}:00"
            except (IndexError, ValueError):
                continue

    return None


def _extraire_modalites_remise(texte: str) -> str | None:
    """Extrait les modalites de remise des offres."""
    patterns = [
        re.compile(
            r"(?:modalit[eé]s?\s*(?:de\s*)?(?:remise|transmission|d[eé]p[oô]t|envoi))"
            r"(.*?)(?=\n\s*(?:article|\d+[\.\)]\s*[A-Z]|ARTICLE))",
            re.IGNORECASE | re.DOTALL,
        ),
        re.compile(
            r"(?:offres?\s*(?:devront?|doivent|seront)\s*[eê]tre\s*(?:remises?|transmises?|d[eé]pos[eé]es?|envoy[eé]es?))"
            r"(.*?)(?:\.|$)",
            re.IGNORECASE,
        ),
    ]

    for pattern in patterns:
        match = pattern.search(texte)
        if match:
            texte_modalite = match.group(1).strip()
            # Nettoyer
            texte_modalite = re.sub(r"\s+", " ", texte_modalite)
            if len(texte_modalite) > 10:
                return texte_modalite[:500]

    # Chercher les plateformes connues
    plateformes = [
        (r"(?:plate[\-\s]?forme|profil\s*d.acheteur)\s*[:\s]*([\w\.\-]+\.(?:fr|com|eu))", "Plateforme"),
        (r"(?:marches[\-\s]?publics|maximilien|aws[\-\s]?achat|place\.marches|klekoon|e[\-\s]?mega|achatpublic)", "Plateforme dematerialisee"),
    ]
    for pattern_str, label in plateformes:
        match = re.search(pattern_str, texte, re.IGNORECASE)
        if match:
            return f"{label}: {match.group(0).strip()}"

    return None


def _extraire_variantes(texte: str) -> bool | None:
    """Detecte si les variantes sont autorisees."""
    if re.search(r"variantes?\s*(?:sont\s*)?autoris[eé]es?", texte, re.IGNORECASE):
        return True
    if re.search(r"variantes?\s*(?:ne\s*sont\s*)?(?:pas\s*)?(?:admises?|autoris[eé]es?|accept[eé]es?)\s*(?:ne|pas|non)", texte, re.IGNORECASE):
        return False
    if re.search(r"(?:pas|non|aucune)\s*(?:de\s*)?variantes?", texte, re.IGNORECASE):
        return False
    if re.search(r"variantes?\s*(?:sont\s*)?(?:admises?|accept[eé]es?)", texte, re.IGNORECASE):
        return True
    return None


def _extraire_lots(texte: str) -> list[dict]:
    """Extrait la description des lots si le marche est alloti."""
    lots = []

    # Pattern "Lot n : description"
    pattern_lot = re.compile(
        r"lot\s*(?:n[o°]?\s*)?(\d+)\s*[\-:\.]\s*(.+?)(?:\n|$)",
        re.IGNORECASE,
    )
    for match in pattern_lot.finditer(texte):
        num = match.group(1)
        desc = match.group(2).strip().rstrip(".")
        if desc and len(desc) > 3:
            lots.append({"numero": int(num), "description": desc[:200]})

    # Deduplication par numero
    seen = set()
    unique = []
    for lot in lots:
        if lot["numero"] not in seen:
            seen.add(lot["numero"])
            unique.append(lot)

    return sorted(unique, key=lambda x: x["numero"])


def _extraire_conditions_participation(texte: str) -> list[str]:
    """Extrait les conditions de participation (CA min, certifications, etc.)."""
    conditions = []

    # Chiffre d'affaires minimum
    match_ca = re.search(
        r"chiffre\s*d.affaires?\s*(?:annuel\s*)?(?:minimum|minimal|sup[eé]rieur|au\s*moins)\s*[àa:]\s*([\d\s\.,]+)\s*(?:EUR|euros?|\u20ac)",
        texte, re.IGNORECASE,
    )
    if match_ca:
        conditions.append(f"Chiffre d'affaires minimum: {match_ca.group(1).strip()} EUR")

    # Certifications requises
    certifs = [
        (r"Qualiopi", "Certification Qualiopi requise"),
        (r"ISO\s*\d+", None),  # capturer le numero
        (r"certification\s+(?:professionnelle\s+)?(?:obligatoire|requise|exig[eé]e)", "Certification professionnelle requise"),
        (r"habilitation", "Habilitation requise"),
    ]
    for pattern_str, label in certifs:
        match = re.search(pattern_str, texte, re.IGNORECASE)
        if match:
            if label:
                conditions.append(label)
            else:
                conditions.append(f"Certification {match.group(0)} requise")

    # Experience minimum
    match_exp = re.search(
        r"(?:exp[eé]rience|r[eé]f[eé]rences?)\s*(?:d.au\s*moins|minimum|de\s*plus\s*de)\s*(\d+)\s*(?:ans?|ann[eé]es?)",
        texte, re.IGNORECASE,
    )
    if match_exp:
        conditions.append(f"Experience minimum: {match_exp.group(1)} ans")

    # Effectif minimum
    match_eff = re.search(
        r"effectif\s*(?:minimum|minimal|d.au\s*moins)\s*[:\s]*(\d+)\s*(?:personnes?|salari[eé]s?|ETP)?",
        texte, re.IGNORECASE,
    )
    if match_eff:
        conditions.append(f"Effectif minimum: {match_eff.group(1)}")

    # Assurance RC Pro
    if re.search(r"assurance\s*(?:responsabilit[eé]\s*civile|RC)\s*(?:professionnelle)?", texte, re.IGNORECASE):
        conditions.append("Assurance RC Professionnelle requise")

    return conditions


def _extraire_duree_marche(texte: str) -> str | None:
    """Extrait la duree prevue du marche."""
    patterns = [
        re.compile(r"dur[eé]e\s*(?:du\s*)?march[eé]\s*[\-:=]\s*(.+?)(?:\.|$)", re.IGNORECASE),
        re.compile(r"dur[eé]e\s*(?:du\s*)?(?:contrat|accord[\-\s]cadre)\s*[\-:=]\s*(.+?)(?:\.|$)", re.IGNORECASE),
        re.compile(r"dur[eé]e\s*(?:pr[eé]vue|estim[eé]e)\s*[\-:=]\s*(.+?)(?:\.|$)", re.IGNORECASE),
        re.compile(r"(\d+)\s*(?:mois|ans?|ann[eé]es?|jours?)\s*(?:renouvelable|reconductible)?", re.IGNORECASE),
    ]

    for pattern in patterns:
        match = pattern.search(texte)
        if match:
            duree = match.group(1).strip() if match.lastindex else match.group(0).strip()
            duree = re.sub(r"\s+", " ", duree)
            if len(duree) > 3:
                return duree[:200]

    return None


def _extraire_montant_estime(texte: str) -> str | None:
    """Extrait le montant estime du marche si mentionne."""
    patterns = [
        re.compile(
            r"(?:montant\s*(?:estim[eé]|pr[eé]visionnel|maximum|global)|"
            r"valeur\s*(?:estim[eé]e|du\s*march[eé])|"
            r"budget\s*(?:pr[eé]visionnel|estim[eé]|allou[eé]|maximum))"
            r"\s*[\-:=]\s*([\d\s\.,]+)\s*(?:EUR|euros?|\u20ac|HT|TTC)",
            re.IGNORECASE,
        ),
        re.compile(
            r"([\d\s\.,]+)\s*(?:EUR|euros?|\u20ac)\s*(?:HT|TTC)\s*(?:estim[eé]|pr[eé]visionnel|maximum)",
            re.IGNORECASE,
        ),
    ]

    for pattern in patterns:
        match = pattern.search(texte)
        if match:
            montant = match.group(1).strip()
            if montant:
                return f"{montant} EUR"

    return None


def extraire_rc(dossier_dce: Path) -> dict:
    """Fonction principale : extrait les informations cles du RC.

    Args:
        dossier_dce: Chemin vers le dossier contenant les fichiers DCE telecharges

    Returns:
        dict avec les informations extraites du RC :
        - fichier_rc: nom du fichier identifie comme RC
        - pieces_exigees: list[str]
        - criteres_attribution: list[dict] avec nom et poids
        - date_limite: str (ISO) ou None
        - modalites_remise: str ou None
        - variantes_autorisees: bool ou None
        - lots: list[dict] avec numero et description
        - conditions_participation: list[str]
        - duree_marche: str ou None
        - montant_estime: str ou None
        - texte_brut: str (les 5000 premiers caracteres du RC)
    """
    resultat = {
        "fichier_rc": None,
        "pieces_exigees": [],
        "criteres_attribution": [],
        "date_limite": None,
        "modalites_remise": None,
        "variantes_autorisees": None,
        "lots": [],
        "conditions_participation": [],
        "duree_marche": None,
        "montant_estime": None,
        "texte_brut": "",
    }

    # Identifier le RC
    fichier_rc = _identifier_rc(dossier_dce)
    if not fichier_rc:
        logger.info("Aucun fichier RC identifie dans le DCE")
        return resultat

    resultat["fichier_rc"] = fichier_rc.name

    # Extraire le texte
    texte = _extraire_texte(fichier_rc)
    if not texte or len(texte.strip()) < 50:
        logger.warning(f"Texte extrait du RC trop court ({len(texte)} chars)")
        return resultat

    logger.info(f"RC extrait: {fichier_rc.name} ({len(texte)} chars)")

    # Stocker un extrait du texte brut
    resultat["texte_brut"] = texte[:5000]

    # Parser chaque information
    resultat["pieces_exigees"] = _extraire_pieces_exigees(texte)
    resultat["criteres_attribution"] = _extraire_criteres_attribution(texte)
    resultat["date_limite"] = _extraire_date_limite(texte)
    resultat["modalites_remise"] = _extraire_modalites_remise(texte)
    resultat["variantes_autorisees"] = _extraire_variantes(texte)
    resultat["lots"] = _extraire_lots(texte)
    resultat["conditions_participation"] = _extraire_conditions_participation(texte)
    resultat["duree_marche"] = _extraire_duree_marche(texte)
    resultat["montant_estime"] = _extraire_montant_estime(texte)

    # Log du resultat
    nb_infos = sum(
        1 for v in resultat.values()
        if v and v is not None and v != [] and v != ""
    )
    logger.info(
        f"RC parse: {nb_infos} infos extraites "
        f"({len(resultat['pieces_exigees'])} pieces, "
        f"{len(resultat['criteres_attribution'])} criteres, "
        f"{len(resultat['lots'])} lots)"
    )

    return resultat


def adapter_dossier(rc_info: dict, ao: dict) -> dict:
    """Analyse les infos du RC et retourne des recommandations d'adaptation.

    Args:
        rc_info: Resultat de extraire_rc()
        ao: Dictionnaire de l'appel d'offres

    Returns:
        dict avec:
        - pieces_manquantes: pieces exigees qu'on ne genere pas encore
        - criteres_focus: criteres les plus ponderes (focus memoire technique)
        - alertes: points d'attention
        - recommandations: liste de recommandations textuelles
    """
    recommandations = []
    alertes = []
    pieces_manquantes = []

    # Pieces qu'on sait generer
    pieces_generees = {
        "memoire technique", "lettre de candidature", "bpu", "bordereau",
        "dpgf", "planning", "cv", "dc1", "dc2", "dume", "acte d'engagement",
        "programme de formation", "references", "moyens techniques",
    }

    # Verifier les pieces exigees vs ce qu'on genere
    for piece in rc_info.get("pieces_exigees", []):
        piece_lower = piece.lower()
        trouvee = any(p in piece_lower for p in pieces_generees)
        if not trouvee:
            pieces_manquantes.append(piece)

    if pieces_manquantes:
        alertes.append(f"{len(pieces_manquantes)} piece(s) exigee(s) non generee(s) automatiquement")
        recommandations.append(
            "Pieces a preparer manuellement : " + ", ".join(pieces_manquantes[:5])
        )

    # Criteres d'attribution : identifier les plus importants pour le memoire
    criteres = rc_info.get("criteres_attribution", [])
    criteres_tries = sorted(criteres, key=lambda c: c.get("poids", 0), reverse=True)
    criteres_focus = criteres_tries[:3]

    if criteres_focus:
        for c in criteres_focus:
            recommandations.append(
                f"Critere '{c['nom']}' ({c['poids']}%) : developper particulierement dans le memoire"
            )

    # Prix vs technique
    prix_poids = 0
    technique_poids = 0
    for c in criteres:
        nom_lower = c.get("nom", "").lower()
        if "prix" in nom_lower or "cout" in nom_lower or "financ" in nom_lower:
            prix_poids += c.get("poids", 0)
        else:
            technique_poids += c.get("poids", 0)

    if prix_poids > 0 and technique_poids > 0:
        if technique_poids >= 60:
            recommandations.append(
                f"Marche a dominante technique ({technique_poids}% vs {prix_poids}% prix) : "
                "privilegier la qualite du memoire technique"
            )
        elif prix_poids >= 60:
            recommandations.append(
                f"Marche a dominante prix ({prix_poids}% vs {technique_poids}% technique) : "
                "optimiser le prix, memoire technique concis et efficace"
            )

    # Conditions de participation
    conditions = rc_info.get("conditions_participation", [])
    for cond in conditions:
        cond_lower = cond.lower()
        if "qualiopi" in cond_lower:
            recommandations.append("Qualiopi exige : inclure le certificat Qualiopi dans le dossier")
        elif "chiffre d'affaires" in cond_lower:
            alertes.append(f"Condition financiere : {cond}")
        elif "experience" in cond_lower or "reference" in cond_lower:
            recommandations.append(f"Condition d'experience : adapter les references clients en consequence")

    # Lots
    lots = rc_info.get("lots", [])
    if lots:
        recommandations.append(
            f"Marche alloti ({len(lots)} lots) : verifier quels lots sont pertinents pour Almera"
        )

    # Date limite
    date_limite = rc_info.get("date_limite")
    if date_limite:
        try:
            dt = datetime.fromisoformat(date_limite)
            jours_restants = (dt - datetime.now()).days
            if jours_restants < 3:
                alertes.append(f"URGENT : date limite dans {jours_restants} jour(s) !")
            elif jours_restants < 7:
                alertes.append(f"Attention : date limite dans {jours_restants} jours")
        except (ValueError, TypeError):
            pass

    # Variantes
    if rc_info.get("variantes_autorisees"):
        recommandations.append("Variantes autorisees : envisager une variante avec e-learning ou certification RS6776")

    return {
        "pieces_manquantes": pieces_manquantes,
        "criteres_focus": criteres_focus,
        "alertes": alertes,
        "recommandations": recommandations,
    }
