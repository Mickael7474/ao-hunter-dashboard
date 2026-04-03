"""
Pre-remplissage des formulaires PDF/Word natifs du DCE (DC1, DC2, DUME).

Detecte automatiquement les formulaires dans le dossier DCE telecharge,
puis les pre-remplit avec les informations Almera (SIRET, adresse, etc.)
en utilisant pypdf pour les PDF (AcroForm) et python-docx pour les DOCX.
"""

import re
import logging
from pathlib import Path

logger = logging.getLogger("ao_hunter.formulaires_natifs")

# Patterns de noms de fichiers pour detecter les formulaires
PATTERNS_FORMULAIRES = [
    re.compile(r"DC\s*1", re.IGNORECASE),
    re.compile(r"DC\s*2", re.IGNORECASE),
    re.compile(r"DUME", re.IGNORECASE),
    re.compile(r"formulaire", re.IGNORECASE),
    re.compile(r"lettre\s*de\s*candidature", re.IGNORECASE),
    re.compile(r"declaration\s*du\s*candidat", re.IGNORECASE),
    re.compile(r"DC\s*4", re.IGNORECASE),
]

# Infos Almera a injecter dans les formulaires
INFOS_ALMERA = {
    "raison_sociale": "AI MENTOR",
    "nom_commercial": "Almera",
    "forme_juridique": "SASU",
    "siret": "98900455100010",
    "siren": "989004551",
    "nda": "11757431975",
    "adresse": "25 rue Campagne Premiere",
    "code_postal": "75014",
    "ville": "Paris",
    "adresse_complete": "25 rue Campagne Premiere, 75014 Paris",
    "pays": "France",
    "representant": "Mickael Bertolla",
    "qualite_representant": "President",
    "representant_complet": "Mickael Bertolla, President",
    "telephone": "+33686680611",
    "tel": "+33686680611",
    "email": "contact@almera.one",
    "site_web": "almera.one",
    "naf": "8559A",
    "code_ape": "8559A",
    "capital": "1 000 EUR",
    "tva_intra": "FR61989004551",
}

# Mapping champs PDF courants -> cle dans INFOS_ALMERA
# Les formulaires DC1/DC2 officiels utilisent des noms de champs varies
MAPPING_CHAMPS_PDF = {
    # Raison sociale / denomination
    "denomination": "raison_sociale",
    "raison_sociale": "raison_sociale",
    "raison sociale": "raison_sociale",
    "nom_entreprise": "raison_sociale",
    "candidat": "raison_sociale",
    "nom du candidat": "raison_sociale",
    "nom candidat": "raison_sociale",
    "societe": "raison_sociale",
    "nom_commercial": "nom_commercial",
    "nom commercial": "nom_commercial",
    "enseigne": "nom_commercial",
    # Forme juridique
    "forme_juridique": "forme_juridique",
    "forme juridique": "forme_juridique",
    "statut_juridique": "forme_juridique",
    "forme": "forme_juridique",
    # SIRET / SIREN
    "siret": "siret",
    "n_siret": "siret",
    "no_siret": "siret",
    "numero_siret": "siret",
    "siren": "siren",
    "n_siren": "siren",
    # NDA
    "nda": "nda",
    "numero_declaration": "nda",
    "declaration_activite": "nda",
    # Adresse
    "adresse": "adresse",
    "adresse_siege": "adresse",
    "siege_social": "adresse_complete",
    "rue": "adresse",
    "code_postal": "code_postal",
    "cp": "code_postal",
    "ville": "ville",
    "commune": "ville",
    "pays": "pays",
    # Representant
    "representant": "representant",
    "nom_representant": "representant",
    "nom representant": "representant",
    "signataire": "representant",
    "nom_signataire": "representant",
    "qualite": "qualite_representant",
    "qualite_signataire": "qualite_representant",
    "fonction": "qualite_representant",
    "qualite_representant": "qualite_representant",
    # Contact
    "telephone": "telephone",
    "tel": "tel",
    "courriel": "email",
    "email": "email",
    "mail": "email",
    "adresse_electronique": "email",
    # Codes
    "naf": "naf",
    "code_naf": "naf",
    "ape": "code_ape",
    "code_ape": "code_ape",
    "capital": "capital",
    "capital_social": "capital",
    "tva": "tva_intra",
    "tva_intra": "tva_intra",
    "tva_intracommunautaire": "tva_intra",
}

# Placeholders texte courants dans les DOCX (entre crochets, pointilles, etc.)
PLACEHOLDERS_DOCX = [
    # Raison sociale
    (re.compile(r"\[?\s*(?:raison\s*sociale|denomination|nom\s*(?:du\s*)?candidat)\s*\]?", re.IGNORECASE), "raison_sociale"),
    (re.compile(r"\.{3,}\s*(?:raison\s*sociale|denomination)", re.IGNORECASE), "raison_sociale"),
    (re.compile(r"(?:raison\s*sociale|denomination)\s*:\s*\.{3,}", re.IGNORECASE), "raison_sociale"),
    # Forme juridique
    (re.compile(r"\[?\s*forme\s*juridique\s*\]?", re.IGNORECASE), "forme_juridique"),
    (re.compile(r"forme\s*juridique\s*:\s*\.{3,}", re.IGNORECASE), "forme_juridique"),
    # SIRET
    (re.compile(r"\[?\s*(?:n[o°]?\s*)?siret\s*\]?", re.IGNORECASE), "siret"),
    (re.compile(r"(?:n[o°]?\s*)?siret\s*:\s*\.{3,}", re.IGNORECASE), "siret"),
    (re.compile(r"\[?\s*(?:n[o°]?\s*)?siren\s*\]?", re.IGNORECASE), "siren"),
    # Adresse
    (re.compile(r"\[?\s*adresse\s*(?:du\s*siege|siege\s*social)?\s*\]?", re.IGNORECASE), "adresse_complete"),
    (re.compile(r"adresse\s*:\s*\.{3,}", re.IGNORECASE), "adresse_complete"),
    (re.compile(r"\[?\s*code\s*postal\s*\]?", re.IGNORECASE), "code_postal"),
    (re.compile(r"\[?\s*(?:ville|commune)\s*\]?", re.IGNORECASE), "ville"),
    # Representant
    (re.compile(r"\[?\s*(?:nom\s*(?:du\s*)?representant|signataire)\s*\]?", re.IGNORECASE), "representant"),
    (re.compile(r"(?:representant|signataire)\s*:\s*\.{3,}", re.IGNORECASE), "representant"),
    (re.compile(r"\[?\s*(?:qualite|fonction)\s*(?:du\s*)?(?:representant|signataire)?\s*\]?", re.IGNORECASE), "qualite_representant"),
    # Contact
    (re.compile(r"\[?\s*(?:telephone|tel)\s*\]?", re.IGNORECASE), "telephone"),
    (re.compile(r"(?:telephone|tel)\s*:\s*\.{3,}", re.IGNORECASE), "telephone"),
    (re.compile(r"\[?\s*(?:courriel|email|mail|adresse\s*electronique)\s*\]?", re.IGNORECASE), "email"),
    (re.compile(r"(?:courriel|email|mail)\s*:\s*\.{3,}", re.IGNORECASE), "email"),
    # NAF / APE
    (re.compile(r"\[?\s*(?:code\s*)?(?:naf|ape)\s*\]?", re.IGNORECASE), "naf"),
    (re.compile(r"(?:code\s*)?(?:naf|ape)\s*:\s*\.{3,}", re.IGNORECASE), "naf"),
    # Capital
    (re.compile(r"\[?\s*capital\s*(?:social)?\s*\]?", re.IGNORECASE), "capital"),
    (re.compile(r"capital\s*(?:social)?\s*:\s*\.{3,}", re.IGNORECASE), "capital"),
    # NDA
    (re.compile(r"\[?\s*(?:n[o°]?\s*)?(?:declaration\s*d.activite|nda)\s*\]?", re.IGNORECASE), "nda"),
    # TVA
    (re.compile(r"\[?\s*(?:n[o°]?\s*)?tva\s*(?:intra(?:communautaire)?)?\s*\]?", re.IGNORECASE), "tva_intra"),
]


def _est_formulaire(fichier: Path) -> bool:
    """Detecte si un fichier est un formulaire DC1/DC2/DUME par son nom."""
    nom = fichier.stem
    for pattern in PATTERNS_FORMULAIRES:
        if pattern.search(nom):
            return True
    return False


def _preremplir_pdf(fichier: Path, sortie: Path) -> bool:
    """Pre-remplit un formulaire PDF (champs AcroForm) avec les infos Almera.

    Returns:
        True si au moins un champ a ete rempli.
    """
    try:
        from pypdf import PdfReader, PdfWriter
    except ImportError:
        logger.warning("pypdf non installe - pre-remplissage PDF impossible")
        return False

    try:
        reader = PdfReader(str(fichier))
    except Exception as e:
        logger.warning(f"Impossible de lire le PDF {fichier.name}: {e}")
        return False

    # Verifier si le PDF a des champs de formulaire
    fields = reader.get_fields()
    if not fields:
        logger.info(f"Pas de champs AcroForm dans {fichier.name}")
        return False

    # Construire le dict de valeurs a remplir
    valeurs = {}
    champs_remplis = 0

    for nom_champ, field_info in fields.items():
        # Normaliser le nom du champ pour le matching
        nom_lower = nom_champ.lower().strip()
        nom_clean = re.sub(r"[\s_\-\.]+", "_", nom_lower)

        # Chercher une correspondance dans le mapping
        cle_almera = None
        for pattern_champ, cle in MAPPING_CHAMPS_PDF.items():
            if pattern_champ in nom_clean or nom_clean in pattern_champ:
                cle_almera = cle
                break

        # Matching plus souple : recherche de sous-chaines
        if not cle_almera:
            for pattern_champ, cle in MAPPING_CHAMPS_PDF.items():
                pattern_clean = re.sub(r"[\s_\-\.]+", "_", pattern_champ)
                if pattern_clean in nom_clean or nom_clean in pattern_clean:
                    cle_almera = cle
                    break

        if cle_almera and cle_almera in INFOS_ALMERA:
            valeurs[nom_champ] = INFOS_ALMERA[cle_almera]
            champs_remplis += 1
            logger.debug(f"  Champ PDF '{nom_champ}' -> {cle_almera} = {INFOS_ALMERA[cle_almera]}")

    if not valeurs:
        logger.info(f"Aucun champ reconnu dans {fichier.name} ({len(fields)} champs)")
        return False

    # Ecrire le PDF pre-rempli
    try:
        writer = PdfWriter()
        writer.append(reader)
        writer.update_page_form_field_values(writer.pages[0], valeurs)

        # Tenter de mettre a jour tous les pages
        for page_num in range(len(writer.pages)):
            try:
                writer.update_page_form_field_values(writer.pages[page_num], valeurs)
            except Exception:
                pass

        with open(str(sortie), "wb") as f:
            writer.write(f)

        logger.info(f"PDF pre-rempli: {fichier.name} ({champs_remplis}/{len(fields)} champs)")
        return True

    except Exception as e:
        logger.error(f"Erreur ecriture PDF {fichier.name}: {e}")
        return False


def _preremplir_docx(fichier: Path, sortie: Path) -> bool:
    """Pre-remplit un formulaire DOCX en remplacant les placeholders.

    Strategies :
    1. Remplir les content controls (structured document tags) si presents
    2. Remplacer les placeholders texte courants (crochets, pointilles)
    3. Remplacer les champs de formulaire Word (form fields)

    Returns:
        True si au moins un remplacement a ete effectue.
    """
    try:
        from docx import Document
    except ImportError:
        logger.warning("python-docx non installe - pre-remplissage DOCX impossible")
        return False

    try:
        doc = Document(str(fichier))
    except Exception as e:
        logger.warning(f"Impossible de lire le DOCX {fichier.name}: {e}")
        return False

    remplacements = 0

    # Strategie 1 : Parcourir les paragraphes et remplacer les placeholders
    for para in doc.paragraphs:
        texte_original = para.text
        if not texte_original.strip():
            continue

        for pattern, cle_almera in PLACEHOLDERS_DOCX:
            if cle_almera in INFOS_ALMERA and pattern.search(texte_original):
                # Remplacer dans les runs pour preserver le formatage
                nouveau_texte = pattern.sub(INFOS_ALMERA[cle_almera], texte_original)
                if nouveau_texte != texte_original:
                    _remplacer_texte_paragraphe(para, texte_original, nouveau_texte)
                    remplacements += 1
                    texte_original = nouveau_texte

    # Strategie 2 : Parcourir les tableaux (tres courant dans les DC1/DC2)
    for table in doc.tables:
        for row in table.rows:
            for i, cell in enumerate(row.cells):
                texte_cell = cell.text.strip()
                if not texte_cell:
                    continue

                # Verifier si la cellule suivante contient un placeholder
                if i + 1 < len(row.cells):
                    cellule_valeur = row.cells[i + 1]
                    texte_valeur = cellule_valeur.text.strip()

                    # Si la cellule courante est un label et la suivante est vide ou avec pointilles
                    if texte_valeur == "" or re.match(r"^\.{3,}$", texte_valeur) or texte_valeur == "...":
                        cle = _detecter_cle_depuis_label(texte_cell)
                        if cle and cle in INFOS_ALMERA:
                            for para in cellule_valeur.paragraphs:
                                if para.runs:
                                    para.runs[0].text = INFOS_ALMERA[cle]
                                    for run in para.runs[1:]:
                                        run.text = ""
                                else:
                                    para.text = INFOS_ALMERA[cle]
                                remplacements += 1

                # Aussi remplacer les placeholders dans les cellules de tableau
                for para in cell.paragraphs:
                    texte_p = para.text
                    if not texte_p.strip():
                        continue
                    for pattern, cle_almera in PLACEHOLDERS_DOCX:
                        if cle_almera in INFOS_ALMERA and pattern.search(texte_p):
                            nouveau = pattern.sub(INFOS_ALMERA[cle_almera], texte_p)
                            if nouveau != texte_p:
                                _remplacer_texte_paragraphe(para, texte_p, nouveau)
                                remplacements += 1
                                texte_p = nouveau

    if remplacements == 0:
        logger.info(f"Aucun placeholder reconnu dans {fichier.name}")
        return False

    try:
        doc.save(str(sortie))
        logger.info(f"DOCX pre-rempli: {fichier.name} ({remplacements} remplacements)")
        return True
    except Exception as e:
        logger.error(f"Erreur sauvegarde DOCX {fichier.name}: {e}")
        return False


def _remplacer_texte_paragraphe(para, ancien: str, nouveau: str):
    """Remplace le texte d'un paragraphe en preservant le formatage du premier run."""
    if not para.runs:
        para.text = nouveau
        return

    # Concatener tous les runs, remplacer, puis remettre dans le premier run
    texte_complet = "".join(run.text for run in para.runs)
    texte_remplace = texte_complet.replace(ancien, nouveau) if ancien in texte_complet else nouveau

    para.runs[0].text = texte_remplace
    for run in para.runs[1:]:
        run.text = ""


def _detecter_cle_depuis_label(label: str) -> str | None:
    """Detecte quelle info Almera correspond a un label de cellule."""
    label_lower = label.lower().strip().rstrip(":").strip()

    mapping_labels = {
        "raison sociale": "raison_sociale",
        "denomination": "raison_sociale",
        "denomination sociale": "raison_sociale",
        "nom commercial": "nom_commercial",
        "enseigne": "nom_commercial",
        "forme juridique": "forme_juridique",
        "statut juridique": "forme_juridique",
        "siret": "siret",
        "n siret": "siret",
        "numero siret": "siret",
        "siren": "siren",
        "adresse": "adresse_complete",
        "adresse du siege": "adresse_complete",
        "siege social": "adresse_complete",
        "code postal": "code_postal",
        "ville": "ville",
        "commune": "ville",
        "pays": "pays",
        "representant": "representant",
        "nom du representant": "representant",
        "signataire": "representant",
        "qualite": "qualite_representant",
        "qualite du representant": "qualite_representant",
        "fonction": "qualite_representant",
        "telephone": "telephone",
        "tel": "telephone",
        "courriel": "email",
        "email": "email",
        "adresse electronique": "email",
        "code naf": "naf",
        "code ape": "code_ape",
        "naf": "naf",
        "ape": "code_ape",
        "capital social": "capital",
        "capital": "capital",
        "tva intracommunautaire": "tva_intra",
        "n tva": "tva_intra",
        "declaration d'activite": "nda",
        "nda": "nda",
        "numero de declaration d'activite": "nda",
    }

    # Normaliser : retirer accents basiques et ponctuation
    label_norm = re.sub(r"[^a-z0-9\s]", "", label_lower)
    label_norm = re.sub(r"\s+", " ", label_norm).strip()

    for pattern, cle in mapping_labels.items():
        if pattern in label_norm or label_norm in pattern:
            return cle

    return None


def preremplir_formulaires(dossier_dce: Path, dossier_sortie: Path) -> list[str]:
    """Detecte et pre-remplit les formulaires DC1/DC2/DUME dans un dossier DCE.

    Args:
        dossier_dce: Chemin vers le dossier contenant les fichiers DCE telecharges
        dossier_sortie: Chemin vers le dossier ou sauvegarder les formulaires pre-remplis

    Returns:
        Liste des noms de fichiers pre-remplis avec succes
    """
    if not dossier_dce.exists() or not dossier_dce.is_dir():
        logger.warning(f"Dossier DCE introuvable: {dossier_dce}")
        return []

    dossier_sortie.mkdir(parents=True, exist_ok=True)

    fichiers_preremplis = []
    extensions_supportees = {".pdf", ".docx", ".doc"}

    # Scanner les fichiers du DCE
    for fichier in sorted(dossier_dce.iterdir()):
        if not fichier.is_file():
            continue
        if fichier.suffix.lower() not in extensions_supportees:
            continue
        if not _est_formulaire(fichier):
            continue

        logger.info(f"Formulaire detecte: {fichier.name}")
        sortie = dossier_sortie / f"prerempli_{fichier.name}"

        try:
            if fichier.suffix.lower() == ".pdf":
                ok = _preremplir_pdf(fichier, sortie)
            elif fichier.suffix.lower() in (".docx", ".doc"):
                ok = _preremplir_docx(fichier, sortie)
            else:
                ok = False

            if ok:
                fichiers_preremplis.append(sortie.name)
                logger.info(f"  -> {sortie.name}")
            else:
                logger.info(f"  -> Pas de champs reconnus dans {fichier.name}")

        except Exception as e:
            logger.error(f"Erreur pre-remplissage {fichier.name}: {e}")

    if fichiers_preremplis:
        logger.info(f"Pre-remplissage: {len(fichiers_preremplis)} formulaire(s) traite(s)")
    else:
        logger.info("Aucun formulaire pre-rempli (pas de formulaires detectes ou pas de champs reconnus)")

    return fichiers_preremplis
