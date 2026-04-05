"""
Export DOCX leger pour Render.
Convertit du Markdown en Word DOCX avec la charte graphique Almera.
Utilise python-docx (pas de dependances C, compatible Render).
"""

import re
import logging
from pathlib import Path
from datetime import datetime

logger = logging.getLogger("ao_hunter.export_docx_render")

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.enum.section import WD_ORIENT
    from docx.oxml.ns import qn, nsdecls
    from docx.oxml import parse_xml
    DOCX_DISPONIBLE = True
except ImportError:
    DOCX_DISPONIBLE = False
    logger.warning("python-docx non installe - export DOCX desactive")

# Couleurs charte Almera
BLEU_FONCE = RGBColor(0x1E, 0x3A, 0x5F) if DOCX_DISPONIBLE else None
BLEU = RGBColor(0x25, 0x63, 0xEB) if DOCX_DISPONIBLE else None
GRIS = RGBColor(0x99, 0x99, 0x99) if DOCX_DISPONIBLE else None
BLANC = RGBColor(0xFF, 0xFF, 0xFF) if DOCX_DISPONIBLE else None
GRIS_CLAIR = RGBColor(0xF3, 0xF4, 0xF6) if DOCX_DISPONIBLE else None

DASHBOARD_DIR = Path(__file__).parent


def _trouver_logo() -> str | None:
    """Cherche le logo Almera dans les emplacements possibles."""
    chemins = [
        DASHBOARD_DIR / "almera_logo.png",
        DASHBOARD_DIR.parent / "almera_logo.png",
    ]
    for chemin in chemins:
        if chemin.exists():
            return str(chemin)
    return None


def _set_font(run, name="Calibri", size=None, bold=False, italic=False, color=None):
    """Applique les proprietes de police a un run."""
    run.font.name = name
    if size:
        run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color


def _ajouter_page_garde(doc, titre_document, sous_titre="", logo_path=None):
    """Ajoute une page de garde avec branding Almera."""
    # Espacement haut
    for _ in range(4):
        doc.add_paragraph("")

    # Logo
    if logo_path:
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(logo_path, width=Inches(2.5))
        except Exception as e:
            logger.warning(f"Impossible d'inserer le logo: {e}")

    doc.add_paragraph("")

    # Titre du document
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(titre_document)
    _set_font(run, size=26, bold=True, color=BLEU_FONCE)

    doc.add_paragraph("")

    # Sous-titre (titre de l'AO)
    if sous_titre:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(sous_titre)
        _set_font(run, size=14, color=BLEU)

    doc.add_paragraph("")

    # Date
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(datetime.now().strftime("%d/%m/%Y"))
    _set_font(run, size=12, color=GRIS)

    doc.add_paragraph("")

    # Mention Almera
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Almera - AI MENTOR | SIRET 98900455100010")
    _set_font(run, size=10, color=GRIS)

    # Saut de page
    doc.add_page_break()


def _ajouter_header_footer(doc):
    """Ajoute header et footer a toutes les sections."""
    for section in doc.sections:
        # Header
        header = section.header
        header.is_linked_to_previous = False
        p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.add_run("Almera - almera.one")
        _set_font(run, size=8, color=GRIS)

        # Footer - page number
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Numero de page via champ Word
        run = p.add_run()
        _set_font(run, size=8, color=GRIS)
        fld_char_begin = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
        run._r.append(fld_char_begin)

        run2 = p.add_run()
        _set_font(run2, size=8, color=GRIS)
        instr_text = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
        run2._r.append(instr_text)

        run3 = p.add_run()
        _set_font(run3, size=8, color=GRIS)
        fld_char_end = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
        run3._r.append(fld_char_end)


def _appliquer_inline(paragraph, texte):
    """Parse le texte inline pour bold/italic et ajoute les runs au paragraph."""
    # Pattern pour **bold**, *italic*, ***bold+italic***
    pattern = re.compile(r'(\*\*\*(.+?)\*\*\*|\*\*(.+?)\*\*|\*(.+?)\*|([^*]+))')
    for match in pattern.finditer(texte):
        if match.group(2):  # ***bold+italic***
            run = paragraph.add_run(match.group(2))
            _set_font(run, size=10.5, bold=True, italic=True)
        elif match.group(3):  # **bold**
            run = paragraph.add_run(match.group(3))
            _set_font(run, size=10.5, bold=True)
        elif match.group(4):  # *italic*
            run = paragraph.add_run(match.group(4))
            _set_font(run, size=10.5, italic=True)
        elif match.group(5):  # texte normal
            run = paragraph.add_run(match.group(5))
            _set_font(run, size=10.5)


def _parser_table(lignes, idx_debut):
    """Parse une table Markdown a partir de la ligne idx_debut.
    Retourne (rows, idx_fin) ou rows est une liste de listes de cellules."""
    rows = []
    i = idx_debut
    while i < len(lignes):
        line = lignes[i].strip()
        if not line.startswith("|"):
            break
        # Ignorer la ligne de separation (|---|---|)
        if re.match(r'^\|[\s\-:|]+\|$', line):
            i += 1
            continue
        cells = [c.strip() for c in line.split("|")[1:-1]]
        if cells:
            rows.append(cells)
        i += 1
    return rows, i


def _ajouter_table(doc, rows):
    """Ajoute un tableau au document avec le style Almera."""
    if not rows:
        return
    nb_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=nb_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j >= nb_cols:
                break
            cell = table.cell(i, j)
            cell.text = ""
            p = cell.paragraphs[0]

            if i == 0:
                # Header row
                shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E3A5F"/>')
                cell._tc.get_or_add_tcPr().append(shading)
                # Retirer les ** du header (tout est deja bold)
                clean_text = cell_text.replace("**", "")
                run = p.add_run(clean_text)
                _set_font(run, size=10, bold=True, color=BLANC)
            else:
                # Alternance de couleurs
                if i % 2 == 0:
                    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F3F4F6"/>')
                    cell._tc.get_or_add_tcPr().append(shading)
                _appliquer_inline(p, cell_text)

    # Bordures legeres
    tbl = table._tbl
    tbl_pr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '  <w:top w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:left w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:right w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="CCCCCC"/>'
        '</w:tblBorders>'
    )
    tbl_pr.append(borders)


def markdown_to_docx(
    contenu_md: str,
    output_path: Path,
    titre_document: str,
    sous_titre: str = "",
    logo_path: str = None,
) -> Path:
    """Convertit un contenu Markdown en fichier DOCX avec branding Almera.

    Args:
        contenu_md: Contenu Markdown a convertir
        output_path: Chemin du fichier DOCX de sortie
        titre_document: Titre affiche sur la page de garde
        sous_titre: Sous-titre (typiquement le titre de l'AO)
        logo_path: Chemin vers le logo (auto-detecte si None)

    Returns:
        Path du fichier genere
    """
    if not DOCX_DISPONIBLE:
        raise ImportError("python-docx n'est pas installe")

    output_path = Path(output_path)
    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # Logo auto-detect
    if logo_path is None:
        logo_path = _trouver_logo()

    # Page de garde
    _ajouter_page_garde(doc, titre_document, sous_titre, logo_path)

    # Header / Footer
    _ajouter_header_footer(doc)

    # Style par defaut du paragraphe
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10.5)
    style.paragraph_format.line_spacing = 1.15

    # Parser le Markdown
    lignes = contenu_md.split("\n")
    i = 0
    while i < len(lignes):
        line = lignes[i]
        stripped = line.strip()

        # Ligne vide
        if not stripped:
            i += 1
            continue

        # Table Markdown
        if stripped.startswith("|") and i + 1 < len(lignes) and lignes[i + 1].strip().startswith("|"):
            rows, i = _parser_table(lignes, i)
            _ajouter_table(doc, rows)
            doc.add_paragraph("")  # Espacement apres table
            continue

        # Titres
        if stripped.startswith("# "):
            level = 0
            text = stripped
            while text.startswith("#"):
                level += 1
                text = text[1:]
            text = text.strip()

            p = doc.add_paragraph()
            run = p.add_run(text)

            if level == 1:
                _set_font(run, size=16, bold=True, color=BLEU_FONCE)
                p.paragraph_format.space_before = Pt(18)
                p.paragraph_format.space_after = Pt(8)
            elif level == 2:
                _set_font(run, size=13, bold=True, color=BLEU)
                p.paragraph_format.space_before = Pt(14)
                p.paragraph_format.space_after = Pt(6)
            elif level >= 3:
                _set_font(run, size=11, bold=True, color=BLEU_FONCE)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(4)

            i += 1
            continue

        # Listes a puces (- ou *)
        if re.match(r'^[\s]*[-*]\s', stripped):
            # Determiner le niveau d'indentation
            indent = len(line) - len(line.lstrip())
            bullet_level = min(indent // 2, 2)
            text = re.sub(r'^[\s]*[-*]\s+', '', line).strip()
            # Checkbox
            text = re.sub(r'^\[[ x]\]\s*', '', text)

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5 + bullet_level * 0.8)
            p.paragraph_format.first_line_indent = Cm(-0.5)

            # Bullet character
            bullet_run = p.add_run("\u2022  ")
            _set_font(bullet_run, size=10.5)

            _appliquer_inline(p, text)
            i += 1
            continue

        # Listes numerotees
        num_match = re.match(r'^(\d+)[.)]\s+(.*)', stripped)
        if num_match:
            num = num_match.group(1)
            text = num_match.group(2)

            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1.5)
            p.paragraph_format.first_line_indent = Cm(-0.5)

            num_run = p.add_run(f"{num}.  ")
            _set_font(num_run, size=10.5, bold=True)

            _appliquer_inline(p, text)
            i += 1
            continue

        # Ligne horizontale (---, ***)
        if re.match(r'^[-*_]{3,}\s*$', stripped):
            # Ajouter un separateur visuel
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run("_" * 50)
            _set_font(run, size=8, color=GRIS)
            i += 1
            continue

        # Paragraphe normal
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.15
        _appliquer_inline(p, stripped)
        i += 1

    # Sauvegarder
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    logger.info(f"DOCX genere: {output_path.name}")
    return output_path
