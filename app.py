"""
AO HUNTER - Dashboard Web
Interface de suivi des appels d'offres, scores et dossiers generes.

Features:
- Dashboard avec stats et graphiques
- Liste AO avec pagination, filtres, recherche live
- Detail AO avec notes, generation dossier
- Export CSV
- Dark mode
- Notifications navigateur (deadline < 3j)
- WebSocket auto-refresh pendant veille
- API REST complete

Usage:
    python dashboard/app.py
    Ouvrir http://localhost:5000
"""

import sys
import os

# Ensure dependencies are on path (project-local .deps folder)
_deps = os.path.join(os.path.dirname(os.path.abspath(__file__)), ".deps")
if os.path.isdir(_deps) and _deps not in sys.path:
    sys.path.insert(0, _deps)

import csv
import json
import io
import threading
import logging
from pathlib import Path
from datetime import datetime, timedelta

sys.path.insert(0, str(Path(__file__).parent.parent))

from flask import (Flask, render_template, jsonify, request,
                   redirect, url_for, Response, make_response)
from flask_socketio import SocketIO
import yaml

app = Flask(__name__)
app.secret_key = os.environ.get("SECRET_KEY", "ao-hunter-2026")
socketio = SocketIO(app, cors_allowed_origins="*")

# --- Scheduler pour veille automatique (Render) ---
_scheduler = None

def init_scheduler():
    """Initialise APScheduler pour la veille automatique toutes les 8h."""
    global _scheduler
    try:
        from apscheduler.schedulers.background import BackgroundScheduler
        from apscheduler.triggers.interval import IntervalTrigger
        from apscheduler.triggers.cron import CronTrigger

        _scheduler = BackgroundScheduler(daemon=True)
        _scheduler.add_job(
            func=_veille_auto,
            trigger=IntervalTrigger(hours=8),
            id="veille_auto",
            name="Veille automatique BOAMP+TED",
            replace_existing=True,
            next_run_time=datetime.now() + timedelta(minutes=2),  # premiere exec 2min apres boot
        )
        _scheduler.add_job(
            func=_rapport_hebdo_auto,
            trigger=CronTrigger(day_of_week="mon", hour=8),
            id="rapport_hebdo",
            name="Rapport hebdomadaire lundi 8h",
            replace_existing=True,
        )
        _scheduler.start()
        logger.info("Scheduler veille auto + rapport hebdo demarre")
    except ImportError:
        logger.warning("APScheduler non installe - veille auto desactivee")
    except Exception as e:
        logger.error(f"Erreur init scheduler: {e}")


def _veille_auto():
    """Fonction executee par le scheduler."""
    try:
        from veille_render import lancer_veille
        result = lancer_veille()
        logger.info(f"Veille auto terminee: {result}")
        socketio.emit("veille_complete", {
            "nouveaux": result["nouveaux"],
            "total": result["total"],
            "auto": True,
        })

        # Alertes instantanees pour les AO a fort potentiel
        if result.get("nouveaux", 0) > 0:
            try:
                from alertes_instantanees import traiter_alertes_batch
                ao_list = charger_ao()
                # Filtrer les AO ajoutes lors de ce cycle (derniere heure)
                from datetime import datetime, timedelta
                seuil = (datetime.now() - timedelta(hours=1)).isoformat()
                nouveaux_ao = [
                    a for a in ao_list
                    if (a.get("date_detection") or a.get("date_ajout") or a.get("date_publication", "")) >= seuil
                ]
                if nouveaux_ao:
                    alertes = traiter_alertes_batch(nouveaux_ao)
                    if alertes:
                        logger.info(f"Alertes instantanees: {len(alertes)} brouillon(s) cree(s)")
                        socketio.emit("alertes_instantanees", {
                            "nb_alertes": len(alertes),
                            "ao_ids": [a.get("ao_id") for a in alertes],
                        })
            except Exception as e:
                logger.error(f"Erreur alertes instantanees: {e}")

    except Exception as e:
        logger.error(f"Erreur veille auto: {e}")

    # Verifier les deadlines et envoyer les rappels
    try:
        from rappels import envoyer_rappels
        appels = charger_ao()
        rappels_result = envoyer_rappels(appels)
        if rappels_result["envoyes"] > 0:
            logger.info(f"Rappels envoyes: {rappels_result['envoyes']}")
    except Exception as e:
        logger.error(f"Erreur rappels: {e}")

    # Veille attributions (AO soumis -> gagne/perdu) + post-mortem auto
    try:
        from veille_attributions import routine_suivi
        attr_result = routine_suivi()
        if attr_result["trouves"] > 0:
            logger.info(
                f"Attributions: {attr_result['trouves']} trouvees "
                f"({attr_result['gagnes']} gagnes, {attr_result.get('post_mortem_declenches', 0)} post-mortem)"
            )
    except Exception as e:
        logger.error(f"Erreur veille attributions: {e}")

    # Veille concurrentielle (1x par jour suffit)
    try:
        from veille_concurrence import lancer_veille_concurrence
        conc = lancer_veille_concurrence()
        logger.info(f"Veille concurrence: {conc['nouvelles']} nouvelles attributions")
    except Exception as e:
        logger.error(f"Erreur veille concurrence: {e}")

    # Benchmark prix (collecte attributions pour stats marche)
    try:
        from benchmark_prix import collecter_attributions
        attribs = collecter_attributions(nb_pages=3)
        logger.info(f"Benchmark prix: {len(attribs)} attributions en base")
    except Exception as e:
        logger.error(f"Erreur benchmark prix: {e}")

    # Veille pre-informations (anticiper les AO 2-6 mois avant)
    try:
        from veille_preinformation import veille_preinfo
        preinfos = veille_preinfo()
        logger.info(f"Veille pre-info: {len(preinfos)} pre-informations")
    except Exception as e:
        logger.error(f"Erreur veille pre-info: {e}")

    # Auto-enrichissement CRM acheteurs
    try:
        from crm_acheteurs import enrichir_acheteur
        appels_crm = charger_ao()
        for ao in appels_crm:
            try:
                enrichir_acheteur(ao)
            except Exception:
                pass
        logger.info(f"CRM: enrichissement de {len(appels_crm)} AO")
    except Exception as e:
        logger.error(f"Erreur enrichissement CRM: {e}")

    # Pipeline full-auto : Go/No-Go -> generation dossier -> brouillon Gmail
    try:
        from pipeline_auto import lancer_pipeline
        pipeline_result = lancer_pipeline()
        if pipeline_result["generes"] > 0:
            logger.info(f"Pipeline auto: {pipeline_result['generes']} dossier(s) genere(s), {pipeline_result['brouillons']} brouillon(s)")
            socketio.emit("pipeline_complete", pipeline_result)
        else:
            logger.info(f"Pipeline auto: aucun dossier genere ({pipeline_result.get('details', [])})")
    except Exception as e:
        logger.error(f"Erreur pipeline auto: {e}")


def _rapport_hebdo_auto():
    """Fonction executee par le scheduler chaque lundi a 8h."""
    try:
        from rapport_hebdo import rapport_et_envoi
        result = rapport_et_envoi()
        if result.get("brouillon_cree"):
            logger.info(f"Rapport hebdo: brouillon cree ({result['rapport'].get('nb_nouveaux_ao', 0)} nouveaux AO)")
        elif result.get("doublon"):
            logger.info("Rapport hebdo: anti-doublon, deja genere cette semaine")
        else:
            logger.warning(f"Rapport hebdo: brouillon non cree ({result.get('erreur', 'inconnu')})")
    except Exception as e:
        logger.error(f"Erreur rapport hebdo auto: {e}")


logger = logging.getLogger("ao_hunter.dashboard")

# Chemins - en local: ao_hunter/resultats/, sur Render: dossier dashboard/
DASHBOARD_DIR = Path(__file__).parent
BASE_DIR = DASHBOARD_DIR.parent
RESULTATS_DIR = BASE_DIR / "resultats"

# Cherche d'abord dans resultats/ (local), sinon dans le dossier dashboard/ (Render)
if (RESULTATS_DIR / "ao_pertinents.json").exists():
    AO_CACHE = RESULTATS_DIR / "ao_pertinents.json"
    NOTES_FILE = RESULTATS_DIR / "ao_notes.json"
else:
    AO_CACHE = DASHBOARD_DIR / "ao_pertinents.json"
    NOTES_FILE = DASHBOARD_DIR / "ao_notes.json"

CONFIG_PATH = BASE_DIR / "config.yaml"

PER_PAGE = 20


# --- Jinja filters ---

@app.template_filter("as_str")
def as_str_filter(value):
    """Convert lists/other types to string for safe display."""
    if isinstance(value, list):
        return ", ".join(str(v) for v in value)
    return str(value) if value else ""


# --- Helpers ---

@app.context_processor
def inject_globals():
    return {"now": datetime.now(), "timedelta": timedelta}


def charger_config() -> dict:
    with open(CONFIG_PATH, "r", encoding="utf-8") as f:
        return yaml.safe_load(f)


def charger_ao() -> list[dict]:
    if not AO_CACHE.exists():
        return []
    with open(AO_CACHE, "r", encoding="utf-8") as f:
        return json.load(f)


def sauvegarder_ao(appels: list[dict]):
    AO_CACHE.write_text(json.dumps(appels, ensure_ascii=False, indent=2), encoding="utf-8")


def charger_notes() -> dict:
    if not NOTES_FILE.exists():
        return {}
    with open(NOTES_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def sauvegarder_notes(notes: dict):
    NOTES_FILE.write_text(json.dumps(notes, ensure_ascii=False, indent=2), encoding="utf-8")


DOSSIERS_INDEX = DASHBOARD_DIR / "dossiers_index.json"
DOSSIERS_GENERES_DIR = DASHBOARD_DIR / "dossiers_generes"
CHECKLIST_FILE = DASHBOARD_DIR / "checklist_etat.json"
COMMENTAIRES_FILE = DASHBOARD_DIR / "commentaires.json"


def _generer_checklist(ao: dict, dossier: dict, modele: dict) -> list[dict]:
    """Genere la checklist de soumission pour un AO."""
    items = [
        {"id": "memoire", "label": "Memoire technique", "categorie": "Offre"},
        {"id": "bpu", "label": "BPU / DPGF (grille tarifaire)", "categorie": "Offre"},
        {"id": "planning", "label": "Planning previsionnel", "categorie": "Offre"},
        {"id": "dc1", "label": "DC1 - Lettre de candidature", "categorie": "Administratif"},
        {"id": "dc2", "label": "DC2 - Declaration du candidat", "categorie": "Administratif"},
        {"id": "kbis", "label": "Extrait Kbis (< 3 mois)", "categorie": "Administratif"},
        {"id": "urssaf", "label": "Attestation URSSAF", "categorie": "Administratif"},
        {"id": "fiscale", "label": "Attestation fiscale", "categorie": "Administratif"},
        {"id": "qualiopi", "label": "Certificat Qualiopi", "categorie": "Administratif"},
        {"id": "rib", "label": "RIB", "categorie": "Administratif"},
        {"id": "assurance", "label": "Attestation RC Pro", "categorie": "Administratif"},
        {"id": "cv", "label": "CV formateurs/consultants", "categorie": "Offre"},
        {"id": "references", "label": "Liste de references", "categorie": "Offre"},
    ]

    # Ajouter les pieces specifiques du modele
    if modele and modele.get("pieces_specifiques"):
        ids_existants = {i["id"] for i in items}
        for p in modele["pieces_specifiques"]:
            pid = p.lower().replace(" ", "_").replace("/", "_")[:20]
            if pid not in ids_existants:
                items.append({"id": pid, "label": p, "categorie": "Specifique"})

    # Marquer auto les pieces du dossier genere
    if dossier and dossier.get("fichiers"):
        fichiers_lower = [f.lower() for f in dossier["fichiers"]]
        for item in items:
            for f in fichiers_lower:
                if item["id"] in f or item["label"].lower().split()[0] in f:
                    item["auto_ok"] = True
                    break

    items.append({"id": "email", "label": "Email de soumission envoye", "categorie": "Soumission"})
    items.append({"id": "depot", "label": "Depot sur la plateforme", "categorie": "Soumission"})

    return items


def _charger_checklist_etat(ao_id: str) -> dict:
    """Charge l'etat de la checklist pour un AO."""
    if CHECKLIST_FILE.exists():
        try:
            data = json.loads(CHECKLIST_FILE.read_text(encoding="utf-8"))
            return data.get(ao_id, {})
        except (json.JSONDecodeError, OSError):
            pass
    return {}


def _sauvegarder_checklist_etat(ao_id: str, etat: dict):
    data = {}
    if CHECKLIST_FILE.exists():
        try:
            data = json.loads(CHECKLIST_FILE.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            data = {}
    data[ao_id] = etat
    CHECKLIST_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _charger_commentaires(ao_id: str) -> list[dict]:
    """Charge les commentaires pour un AO."""
    if COMMENTAIRES_FILE.exists():
        try:
            data = json.loads(COMMENTAIRES_FILE.read_text(encoding="utf-8"))
            return data.get(ao_id, [])
        except (json.JSONDecodeError, OSError):
            pass
    return []


def _sauvegarder_commentaire(ao_id: str, commentaire: dict):
    data = {}
    if COMMENTAIRES_FILE.exists():
        try:
            data = json.loads(COMMENTAIRES_FILE.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            data = {}
    if ao_id not in data:
        data[ao_id] = []
    data[ao_id].append(commentaire)
    COMMENTAIRES_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


def _scan_dossiers_dir(base_dir: Path) -> list[dict]:
    """Scanne un repertoire pour lister les dossiers generes."""
    dossiers = []
    if not base_dir.exists():
        return dossiers
    for d in sorted(base_dir.iterdir(), reverse=True):
        if d.is_dir() and d.name not in ("__pycache__", "test_dc1_dc2"):
            fichiers = [f.name for f in d.glob("*") if f.is_file()]
            if fichiers:
                dossiers.append({
                    "nom": d.name,
                    "chemin": str(d),
                    "nb_fichiers": len(fichiers),
                    "fichiers": fichiers,
                    "date_creation": datetime.fromtimestamp(d.stat().st_ctime).strftime("%Y-%m-%d %H:%M"),
                })
    return dossiers


def lister_dossiers() -> list[dict]:
    """Liste tous les dossiers en fusionnant filesystem + index JSON."""
    # Scanner les repertoires physiques
    dossiers = _scan_dossiers_dir(RESULTATS_DIR)
    dossiers.extend(_scan_dossiers_dir(DOSSIERS_GENERES_DIR))

    # Fusionner avec dossiers_index.json (contient les dossiers perdus apres redeploy)
    if DOSSIERS_INDEX.exists():
        try:
            with open(DOSSIERS_INDEX, "r", encoding="utf-8") as f:
                index_data = json.load(f)
            if isinstance(index_data, list):
                noms_existants = {d["nom"] for d in dossiers}
                for entry in index_data:
                    nom = entry.get("nom", "")
                    if nom and nom not in noms_existants:
                        # Normaliser la structure pour matcher _scan_dossiers_dir
                        dossiers.append({
                            "nom": nom,
                            "chemin": entry.get("chemin", ""),
                            "nb_fichiers": entry.get("nb_fichiers", 0),
                            "fichiers": entry.get("fichiers", []),
                            "date_creation": entry.get("date_creation", ""),
                            "ao_id": entry.get("ao_id", ""),
                            "ao_titre": entry.get("ao_titre", ""),
                            "source": "index",
                        })
        except (json.JSONDecodeError, OSError):
            pass

    return dossiers


def stats_ao(appels: list[dict]) -> dict:
    now = datetime.now()
    stats = {
        "total": len(appels),
        "sources": {},
        "score_moyen": 0,
        "avec_deadline": 0,
        "deadline_proche": 0,
        "expires": 0,
        "par_statut": {},
        "scores_distribution": {"0-20": 0, "20-40": 0, "40-60": 0, "60-80": 0, "80-100": 0},
        "regions": {},
    }
    scores = []
    for ao in appels:
        src = ao.get("source", "Inconnu")
        stats["sources"][src] = stats["sources"].get(src, 0) + 1

        score = ao.get("score_pertinence", 0) or 0
        if score > 0:
            scores.append(score)
        pct = int(score * 100)
        if pct >= 80: stats["scores_distribution"]["80-100"] += 1
        elif pct >= 60: stats["scores_distribution"]["60-80"] += 1
        elif pct >= 40: stats["scores_distribution"]["40-60"] += 1
        elif pct >= 20: stats["scores_distribution"]["20-40"] += 1
        else: stats["scores_distribution"]["0-20"] += 1

        statut = ao.get("statut", "nouveau")
        stats["par_statut"][statut] = stats["par_statut"].get(statut, 0) + 1

        region_raw = ao.get("region") or ""
        if isinstance(region_raw, list):
            region_raw = ", ".join(str(r) for r in region_raw) if region_raw else ""
        region = str(region_raw)[:20]
        if region:
            stats["regions"][region] = stats["regions"].get(region, 0) + 1

        dl = ao.get("date_limite", "")
        if dl:
            stats["avec_deadline"] += 1
            try:
                date_dl = datetime.fromisoformat(dl.split("T")[0])
                delta = (date_dl - now).days
                if delta < 0:
                    stats["expires"] += 1
                elif delta <= 7:
                    stats["deadline_proche"] += 1
            except (ValueError, TypeError):
                pass

    stats["score_moyen"] = sum(scores) / len(scores) if scores else 0
    # Top 10 regions
    stats["top_regions"] = dict(sorted(stats["regions"].items(), key=lambda x: x[1], reverse=True)[:10])
    return stats


def filtrer_ao(appels, source="", score_min=0, statut="", recherche="", tri="score"):
    filtre = appels
    if source:
        filtre = [a for a in filtre if a.get("source") == source]
    if score_min > 0:
        filtre = [a for a in filtre if (a.get("score_pertinence") or 0) >= score_min]
    if statut:
        filtre = [a for a in filtre if a.get("statut") == statut]
    if recherche:
        q = recherche.lower()
        filtre = [a for a in filtre if q in (a.get("titre") or "").lower()
                  or q in (a.get("description") or "").lower()
                  or q in (a.get("acheteur") or "").lower()]
    if tri == "score":
        filtre.sort(key=lambda a: a.get("score_pertinence") or 0, reverse=True)
    elif tri == "date":
        filtre.sort(key=lambda a: a.get("date_limite") or "", reverse=True)
    elif tri == "titre":
        filtre.sort(key=lambda a: (a.get("titre") or "").lower())
    return filtre


def _ao_urgents(appels: list[dict], max_days=7) -> list[dict]:
    now = datetime.now()
    urgents = []
    for ao in appels:
        dl = ao.get("date_limite", "")
        if not dl:
            continue
        try:
            date_dl = datetime.fromisoformat(dl.split("T")[0])
            delta = (date_dl - now).days
            if 0 <= delta <= max_days:
                ao_copy = dict(ao)
                ao_copy["jours_restants"] = delta
                urgents.append(ao_copy)
        except (ValueError, TypeError):
            pass
    urgents.sort(key=lambda a: a.get("jours_restants", 99))
    return urgents


STATUTS_KANBAN = ["nouveau", "analyse", "candidature", "soumis", "gagne", "perdu", "ignore"]


def _calculer_top_opportunites(appels: list[dict], top_n: int = 5) -> list[dict]:
    """Calcule les top AO les plus rentables a traiter maintenant.

    Score d'opportunite base sur : pertinence, budget zone confort (5k-100k),
    concurrence faible, deadline 7-45 jours, Go/No-Go, signaux faibles bas,
    accessibilite elevee. Uniquement statut nouveau/analyse.
    Pas d'appel API, tout est calcul local.
    """
    from estimation_marche import estimer_marche

    now = datetime.now()
    candidats = []

    for ao in appels:
        # Filtre statut : uniquement nouveau ou analyse
        statut = (ao.get("statut") or "nouveau").lower()
        if statut not in ("nouveau", "analyse"):
            continue

        # Filtre deadline : entre 7 et 45 jours
        dl = ao.get("date_limite", "")
        if not dl:
            continue
        try:
            date_dl = datetime.fromisoformat(dl.split("T")[0])
            jours_restants = (date_dl - now).days
        except (ValueError, TypeError):
            continue
        if jours_restants < 7 or jours_restants > 45:
            continue

        # Estimation marche (mode light = pas d'appel API DECP)
        try:
            estimation = estimer_marche(ao, light=True)
        except Exception:
            continue

        budget_montant = estimation["budget"].get("montant", 0)
        concurrence_niveau = estimation["concurrence"].get("niveau", "moyen")
        concurrence_score = estimation["concurrence"].get("score", 50)
        accessibilite_score = estimation["accessibilite"].get("score", 50)
        score_pertinence = ao.get("score_pertinence") or 0

        # --- Calcul du score d'opportunite (0-100) ---
        score_opp = 0

        # 1. Pertinence (25 pts max)
        score_opp += score_pertinence * 25

        # 2. Budget dans la zone confort 5k-100k (20 pts max)
        if 5_000 <= budget_montant <= 100_000:
            # Maximum au sweet spot 15k-50k
            if 15_000 <= budget_montant <= 50_000:
                score_opp += 20
            elif 10_000 <= budget_montant <= 70_000:
                score_opp += 15
            else:
                score_opp += 10
        elif budget_montant < 5_000:
            score_opp += 2
        else:
            score_opp += 5

        # 3. Concurrence faible (20 pts max)
        # concurrence_score: 0=pas de concurrence, 95=tres forte
        score_opp += max(0, (100 - concurrence_score) * 0.2)

        # 4. Deadline ideale (10 pts max)
        # Sweet spot : 10-30 jours
        if 10 <= jours_restants <= 30:
            score_opp += 10
        elif 7 <= jours_restants < 10:
            score_opp += 6
        else:  # 30-45 jours
            score_opp += 7

        # 5. Accessibilite (15 pts max)
        score_opp += accessibilite_score * 0.15

        # 6. Signaux faibles bas (10 pts max)
        try:
            from signaux_faibles import detecter_signaux
            signaux = detecter_signaux(ao)
            risque = signaux.get("score_risque", 0)
            score_opp += max(0, (100 - risque) * 0.1)
        except (ImportError, Exception):
            score_opp += 5  # Bonus neutre si pas dispo

        score_opp = min(100, max(0, round(score_opp, 1)))

        candidats.append({
            "ao": ao,
            "score_opportunite": score_opp,
            "budget_estime": budget_montant,
            "budget_confiance": estimation["budget"].get("confiance", "faible"),
            "concurrence_niveau": concurrence_niveau,
            "accessibilite": accessibilite_score,
            "jours_restants": jours_restants,
            "date_limite": dl,
        })

    # Trier par score d'opportunite decroissant
    candidats.sort(key=lambda c: c["score_opportunite"], reverse=True)
    return candidats[:top_n]

PRESTATIONS_KEYWORDS = {
    "Formation": ["formation", "formateur", "pedagogie", "stagiaire", "apprenant",
                  "competences", "certification", "qualiopi", "cpf", "opco",
                  "apprentissage", "enseignement", "e-learning", "module",
                  "programme pedagogique", "parcours de formation"],
    "Consulting / AMO": ["conseil", "consulting", "accompagnement", "audit",
                         "assistance a maitrise", "amo", "diagnostic", "expertise",
                         "preconisation", "etude", "strategie", "transformation",
                         "conduite du changement", "schema directeur"],
    "Developpement": ["developpement", "logiciel", "application", "site web",
                      "plateforme", "api", "integration", "maintenance applicative",
                      "tma", "devops", "cloud", "hebergement", "infrastructure",
                      "systeme d'information", "numerique", "digital"],
}


def detecter_prestations(ao: dict) -> list[dict]:
    """Detecte les types de prestations pertinentes pour un AO."""
    texte = f"{ao.get('titre', '')} {ao.get('description', '')}".lower()
    resultats = []
    for presta, mots in PRESTATIONS_KEYWORDS.items():
        matches = [m for m in mots if m in texte]
        if matches:
            score = min(100, len(matches) * 25)
            resultats.append({"type": presta, "score": score, "mots_cles": matches[:5]})
    resultats.sort(key=lambda x: x["score"], reverse=True)
    return resultats


# --- Pages ---

@app.route("/")
def index():
    appels = charger_ao()
    dossiers = lister_dossiers()
    statistiques = stats_ao(appels)

    # Top opportunites du jour
    try:
        top_opportunites = _calculer_top_opportunites(appels)
    except Exception as e:
        logger.warning(f"Erreur calcul top opportunites: {e}")
        top_opportunites = []

    return render_template(
        "index.html",
        stats=statistiques,
        nb_dossiers=len(dossiers),
        top_ao=sorted(appels, key=lambda a: a.get("score_pertinence") or 0, reverse=True)[:10],
        urgents=_ao_urgents(appels),
        top_opportunites=top_opportunites,
    )


@app.route("/ao")
def liste_ao():
    appels = charger_ao()
    source = request.args.get("source", "")
    score_min_str = request.args.get("score_min", "0")
    statut = request.args.get("statut", "")
    recherche = request.args.get("q", "")
    tri = request.args.get("tri", "score")
    page = request.args.get("page", "1")

    try:
        score_min = float(score_min_str)
    except ValueError:
        score_min = 0
    try:
        page = max(1, int(page))
    except ValueError:
        page = 1

    filtre = filtrer_ao(appels, source, score_min, statut, recherche, tri)

    total_pages = max(1, (len(filtre) + PER_PAGE - 1) // PER_PAGE)
    page = min(page, total_pages)
    page_appels = filtre[(page - 1) * PER_PAGE: page * PER_PAGE]

    sources = sorted(set(a.get("source", "") for a in appels))
    statuts = sorted(set(a.get("statut", "") for a in appels))

    # Pre-calculer estimations pour la page courante (light=True, pas d'API DECP)
    estimations = {}
    try:
        from estimation_marche import estimer_marche
        for ao in page_appels:
            try:
                estimations[ao.get("id")] = estimer_marche(ao, light=True)
            except Exception:
                pass
    except ImportError:
        pass

    return render_template(
        "ao_liste.html",
        appels=page_appels,
        estimations=estimations,
        total=len(appels),
        total_filtre=len(filtre),
        sources=sources,
        statuts=statuts,
        filtre_source=source,
        filtre_score_min=score_min,
        filtre_statut=statut,
        filtre_recherche=recherche,
        filtre_tri=tri,
        page=page,
        total_pages=total_pages,
        per_page=PER_PAGE,
    )


def _calculer_roi_stats():
    """Calcule toutes les stats ROI a partir des donnees. Retourne un dict."""
    appels = charger_ao()
    dossiers = lister_dossiers()

    # Stats pipeline
    nb_total = len(appels)
    nb_analyses = len([a for a in appels if a.get("statut") not in ("nouveau", "ignore")])
    nb_candidatures = len([a for a in appels if a.get("statut") in ("candidature", "soumis", "gagne", "perdu")])
    nb_soumis = len([a for a in appels if a.get("statut") in ("soumis", "gagne", "perdu")])
    nb_gagnes = len([a for a in appels if a.get("statut") == "gagne"])
    nb_perdus = len([a for a in appels if a.get("statut") == "perdu"])
    nb_ignores = len([a for a in appels if a.get("statut") == "ignore"])
    nb_dossiers = len(dossiers)

    # Taux de conversion
    taux_conversion = (nb_gagnes / nb_soumis * 100) if nb_soumis > 0 else 0

    # Estimation temps economise (en heures)
    temps_veille = nb_total * 0.5
    temps_analyse = nb_analyses * 0.5
    temps_dossiers = nb_dossiers * 8
    temps_total_h = temps_veille + temps_analyse + temps_dossiers

    # Cout API estime
    cout_scoring = nb_total * 0.003
    cout_par_dossier = 0.15
    cout_generation = nb_dossiers * cout_par_dossier
    cout_api_total = cout_scoring + cout_generation

    # Valeur des contrats
    valeur_gagnes = sum(a.get("budget_estime") or 0 for a in appels if a.get("statut") == "gagne")
    valeur_soumis = sum(a.get("budget_estime") or 0 for a in appels if a.get("statut") == "soumis")
    valeur_pipeline = sum(a.get("budget_estime") or 0 for a in appels
                         if a.get("statut") in ("analyse", "candidature"))

    # CA en cours = budget soumis x taux de conversion historique
    ca_en_cours = valeur_soumis * (taux_conversion / 100) if taux_conversion > 0 else 0

    # Temps moyen par dossier (basé sur timestamps de generation)
    temps_moyen_dossier = 0
    if dossiers:
        durees = []
        for d in dossiers:
            dc = d.get("date_creation", "")
            if dc:
                durees.append(1)  # chaque dossier ~ quelques minutes en auto
        temps_moyen_dossier = round(temps_dossiers * 60 / len(dossiers), 1) if dossiers else 0  # min economisees/dossier

    # ROI
    tarif_horaire = 62.5  # 500 EUR/jour / 8h
    cout_temps_valorise = temps_total_h * tarif_horaire
    cout_total = cout_api_total + (nb_dossiers * 5)
    roi_ratio = (valeur_gagnes / cout_total) if cout_total > 0 else 0
    roi_global = (valeur_gagnes / (cout_api_total + cout_temps_valorise)) if (cout_api_total + cout_temps_valorise) > 0 else 0

    tarif_consultant = 500
    jours_economises = temps_total_h / 8
    economie_consultant = jours_economises * tarif_consultant

    # Stats par source avec taux conversion
    sources_stats = {}
    for ao in appels:
        src = ao.get("source", "Inconnu")
        if src not in sources_stats:
            sources_stats[src] = {"total": 0, "soumis": 0, "gagnes": 0, "perdus": 0, "valeur": 0, "taux": 0}
        sources_stats[src]["total"] += 1
        if ao.get("statut") in ("soumis", "gagne", "perdu"):
            sources_stats[src]["soumis"] += 1
        if ao.get("statut") == "gagne":
            sources_stats[src]["gagnes"] += 1
            sources_stats[src]["valeur"] += ao.get("budget_estime") or 0
        if ao.get("statut") == "perdu":
            sources_stats[src]["perdus"] += 1
    for src, s in sources_stats.items():
        s["taux"] = round(s["gagnes"] / s["soumis"] * 100, 1) if s["soumis"] > 0 else 0

    # Evolution mensuelle (AO detectes et gagnes par mois)
    evolution_mensuelle = {}
    for ao in appels:
        dp = ao.get("date_publication", "")
        if dp and len(dp) >= 7:
            mois = dp[:7]  # YYYY-MM
            if mois not in evolution_mensuelle:
                evolution_mensuelle[mois] = {"detectes": 0, "gagnes": 0}
            evolution_mensuelle[mois]["detectes"] += 1
            if ao.get("statut") == "gagne":
                evolution_mensuelle[mois]["gagnes"] += 1
    # Trier par mois
    evolution_mensuelle = dict(sorted(evolution_mensuelle.items()))

    # Top acheteurs
    acheteurs_stats = {}
    for ao in appels:
        acheteur = ao.get("acheteur", "Inconnu")
        if not acheteur:
            acheteur = "Inconnu"
        if acheteur not in acheteurs_stats:
            acheteurs_stats[acheteur] = {"total": 0, "soumis": 0, "gagnes": 0, "valeur": 0}
        acheteurs_stats[acheteur]["total"] += 1
        if ao.get("statut") in ("soumis", "gagne", "perdu"):
            acheteurs_stats[acheteur]["soumis"] += 1
        if ao.get("statut") == "gagne":
            acheteurs_stats[acheteur]["gagnes"] += 1
            acheteurs_stats[acheteur]["valeur"] += ao.get("budget_estime") or 0
    # Top 10 par nombre d'interactions
    top_acheteurs = sorted(acheteurs_stats.items(), key=lambda x: x[1]["total"], reverse=True)[:10]

    # Score moyen par statut
    scores_par_statut = {}
    for ao in appels:
        statut = ao.get("statut", "nouveau")
        score = ao.get("score_pertinence")
        if score is not None:
            if statut not in scores_par_statut:
                scores_par_statut[statut] = {"somme": 0, "count": 0}
            scores_par_statut[statut]["somme"] += score
            scores_par_statut[statut]["count"] += 1
    scores_moyens = {}
    for statut, data in scores_par_statut.items():
        scores_moyens[statut] = round(data["somme"] / data["count"] * 100, 1) if data["count"] > 0 else 0

    # Funnel data (pour les barres CSS)
    funnel_max = max(nb_total, 1)
    funnel = [
        {"label": "Detectes", "count": nb_total, "pct": 100},
        {"label": "Analyses", "count": nb_analyses, "pct": round(nb_analyses / funnel_max * 100, 1)},
        {"label": "Candidatures", "count": nb_candidatures, "pct": round(nb_candidatures / funnel_max * 100, 1)},
        {"label": "Soumis", "count": nb_soumis, "pct": round(nb_soumis / funnel_max * 100, 1)},
        {"label": "Gagnes", "count": nb_gagnes, "pct": round(nb_gagnes / funnel_max * 100, 1)},
    ]

    return {
        "nb_total": nb_total, "nb_analyses": nb_analyses, "nb_candidatures": nb_candidatures,
        "nb_soumis": nb_soumis, "nb_gagnes": nb_gagnes, "nb_perdus": nb_perdus,
        "nb_ignores": nb_ignores, "nb_dossiers": nb_dossiers,
        "taux_conversion": round(taux_conversion, 1),
        "temps_veille": temps_veille, "temps_analyse": temps_analyse,
        "temps_dossiers": temps_dossiers, "temps_total_h": temps_total_h,
        "temps_moyen_dossier": temps_moyen_dossier,
        "cout_scoring": round(cout_scoring, 2), "cout_generation": round(cout_generation, 2),
        "cout_api_total": round(cout_api_total, 2), "cout_par_dossier": cout_par_dossier,
        "valeur_gagnes": valeur_gagnes, "valeur_soumis": valeur_soumis,
        "valeur_pipeline": valeur_pipeline, "ca_en_cours": ca_en_cours,
        "roi_ratio": round(roi_ratio, 1), "roi_global": round(roi_global, 1),
        "economie_consultant": economie_consultant,
        "jours_economises": jours_economises,
        "sources_stats": sources_stats,
        "evolution_mensuelle": evolution_mensuelle,
        "top_acheteurs": top_acheteurs,
        "scores_moyens": scores_moyens,
        "funnel": funnel,
    }


@app.route("/autofill")
def autofill_page():
    """Page d'aide auto-remplissage plateformes de depot (bookmarklet + userscript)."""
    return render_template("autofill.html")


@app.route("/roi")
def roi():
    """Tableau de bord ROI ameliore - stats avancees, funnel, evolution."""
    stats = _calculer_roi_stats()
    return render_template("roi.html", **stats)


@app.route("/statistiques")
def page_statistiques():
    """Page statistiques avancees (funnel, win rate, sources, acheteurs, ROI)."""
    return render_template("statistiques.html")


@app.route("/upload-gagnant", methods=["GET", "POST"])
def page_upload_gagnant():
    """Page pour importer un dossier gagnant (fichiers .md/.docx/.pdf)."""
    if request.method == "POST":
        import shutil
        from memoire_adaptative import sauvegarder_memoire_gagnant

        titre = request.form.get("titre", "Dossier gagnant")
        acheteur = request.form.get("acheteur", "")
        fichiers = request.files.getlist("fichiers")

        if not fichiers:
            return render_template("upload_gagnant.html", erreur="Aucun fichier selectionne")

        # Creer un dossier pour stocker les fichiers
        clean_titre = re.sub(r"[^a-zA-Z0-9_-]", "_", titre)[:50]
        dossier_nom = f"GAGNANT_{clean_titre}_{datetime.now().strftime('%Y%m%d')}"
        dossier_path = DOSSIERS_DIR / dossier_nom
        dossier_path.mkdir(parents=True, exist_ok=True)

        fichiers_sauves = []
        for f in fichiers:
            if f.filename:
                safe_name = re.sub(r"[^a-zA-Z0-9_.-]", "_", f.filename)
                f.save(str(dossier_path / safe_name))
                fichiers_sauves.append(safe_name)

        # Indexer comme memoire gagnant
        ao_fictif = {"id": dossier_nom, "titre": titre, "acheteur": acheteur, "statut": "gagne"}
        try:
            sauvegarder_memoire_gagnant(ao_fictif)
        except Exception as e:
            logger.warning(f"Erreur indexation memoire gagnant: {e}")

        return render_template("upload_gagnant.html",
                               succes=f"{len(fichiers_sauves)} fichier(s) importe(s) dans {dossier_nom}",
                               fichiers=fichiers_sauves)

    return render_template("upload_gagnant.html")


@app.route("/brouillons")
def page_brouillons():
    """Page listant les brouillons Gmail (pipeline auto)."""
    filtre_ao = request.args.get("ao_hunter", "1") == "1"
    try:
        from brouillons_gmail import lister_brouillons
        brouillons = lister_brouillons(max_results=50, filtre_ao_hunter=filtre_ao)
    except Exception as e:
        logger.error(f"Erreur lecture brouillons: {e}")
        brouillons = []
    return render_template("brouillons.html", brouillons=brouillons, filtre_ao=filtre_ao)


@app.route("/api/brouillons")
def api_brouillons():
    """GET /api/brouillons - Liste des brouillons Gmail en JSON."""
    filtre_ao = request.args.get("ao_hunter", "0") == "1"
    try:
        from brouillons_gmail import lister_brouillons
        return jsonify(lister_brouillons(max_results=50, filtre_ao_hunter=filtre_ao))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/kanban")
def kanban():
    appels = charger_ao()
    colonnes = {}
    for s in STATUTS_KANBAN:
        colonnes[s] = sorted(
            [a for a in appels if a.get("statut", "nouveau") == s],
            key=lambda a: a.get("score_pertinence") or 0, reverse=True
        )[:50]  # max 50 par colonne
    # Stats conversion
    total_soumis = len([a for a in appels if a.get("statut") in ("soumis", "gagne", "perdu")])
    gagnes = len([a for a in appels if a.get("statut") == "gagne"])
    taux_conversion = (gagnes / total_soumis * 100) if total_soumis > 0 else 0
    sources = sorted(set(a.get("source", "") for a in appels if a.get("source")))

    # Pre-calculer estimations pour le kanban (light=True, pas d'API DECP)
    estimations_kanban = {}
    try:
        from estimation_marche import estimer_marche
        for ao in appels:
            try:
                estimations_kanban[ao.get("id")] = estimer_marche(ao, light=True)
            except Exception:
                pass
    except ImportError:
        pass

    return render_template("kanban.html", colonnes=colonnes, statuts=STATUTS_KANBAN,
                           taux_conversion=taux_conversion, total_soumis=total_soumis,
                           gagnes=gagnes, sources=sources, estimations=estimations_kanban)


@app.route("/ao/<path:ao_id>/statut-ajax", methods=["POST"])
def changer_statut_ajax(ao_id):
    """Change le statut via AJAX (pour le drag & drop Kanban)."""
    appels = charger_ao()
    data = request.get_json()
    nouveau_statut = data.get("statut", "nouveau")
    for ao in appels:
        if ao.get("id") == ao_id:
            ao["statut"] = nouveau_statut
            sauvegarder_ao(appels)
            # Memoire adaptative : indexer le memoire si AO gagne
            if nouveau_statut == "gagne":
                try:
                    from memoire_adaptative import sauvegarder_memoire_gagnant
                    sauvegarder_memoire_gagnant(ao)
                except Exception as e:
                    logger.warning(f"Memoire adaptative: erreur sauvegarde pour {ao_id}: {e}")
            # Recalibrer le scoring predictif si gagne ou perdu
            if nouveau_statut in ("gagne", "perdu"):
                try:
                    from scoring_predictif import calibrer_auto
                    calibrer_auto()
                except Exception as e:
                    logger.warning(f"Scoring predictif: erreur calibration pour {ao_id}: {e}")
            return jsonify({"status": "ok", "id": ao_id, "statut": nouveau_statut})
    return jsonify({"error": "AO non trouve"}), 404


@app.route("/ao/<path:ao_id>")
def detail_ao(ao_id):
    appels = charger_ao()
    ao = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao:
        return redirect(url_for("liste_ao"))

    # Dossier genere - chercher dans resultats/ et dossiers_generes/
    dossier_genere = None
    clean_id = ao_id.replace("BOAMP-", "").replace("PLACE-", "").replace("TED-", "").replace("MSEC-", "").replace("AWS-", "")
    for base in [RESULTATS_DIR, DOSSIERS_GENERES_DIR]:
        if not base.exists():
            continue
        for d in base.iterdir():
            if d.is_dir() and clean_id in d.name:
                dossier_genere = {
                    "nom": d.name,
                    "fichiers": sorted([f.name for f in d.glob("*") if f.is_file()]),
                }
                break
        if dossier_genere:
            break
    # Fallback: chercher dans dossiers_index.json
    if not dossier_genere and DOSSIERS_INDEX.exists():
        try:
            idx = json.loads(DOSSIERS_INDEX.read_text(encoding="utf-8"))
            if isinstance(idx, list):
                for entry in idx:
                    if entry.get("ao_id") == ao_id or (clean_id and clean_id in entry.get("nom", "")):
                        dossier_genere = {
                            "nom": entry.get("nom", ""),
                            "fichiers": sorted(entry.get("fichiers", [])),
                            "source": "index",
                        }
                        break
        except (json.JSONDecodeError, OSError):
            pass

    # Notes
    notes = charger_notes()
    note = notes.get(ao_id, "")

    # Detection lots si pas encore fait
    if "lots_detectes" not in ao:
        from veille_render import detecter_lots_pertinents
        lots = detecter_lots_pertinents(ao)
        if lots:
            ao["lots_detectes"] = lots
            ao["lots_pertinents"] = sum(1 for l in lots if l.get("pertinent"))

    prestations = detecter_prestations(ao)

    # Go/No-Go rapide
    from analyse_dce import go_no_go
    go_nogo = go_no_go(ao)

    # Modele de reponse recommande
    from modeles_reponse import detecter_type_prestation, get_modele
    type_presta = detecter_type_prestation(ao)
    modele = get_modele(type_presta)

    # Checklist de soumission
    checklist = _generer_checklist(ao, dossier_genere, modele)
    checklist_etat = _charger_checklist_etat(ao_id)

    # Commentaires
    commentaires = _charger_commentaires(ao_id)

    # Estimation budget / concurrence / accessibilite
    try:
        from estimation_marche import estimer_marche
        estimation = estimer_marche(ao)
    except Exception:
        estimation = None

    # Données DECP réelles
    try:
        from decp_data import rechercher_marches_similaires
        decp = rechercher_marches_similaires(ao)
    except Exception:
        decp = None

    # Score acheteur
    try:
        from score_acheteur import scorer_acheteur
        score_acheteur_data = scorer_acheteur(ao)
    except Exception:
        score_acheteur_data = None

    # AO similaires (si gagne)
    ao_similaires = []
    if ao.get("statut") == "gagne":
        try:
            from alertes_similaires import trouver_similaires
            ao_similaires = trouver_similaires(ao, appels, n=5)
        except Exception:
            ao_similaires = []

    # Post-mortem (si perdu)
    post_mortem_data = None
    if ao.get("statut") == "perdu":
        try:
            from post_mortem import analyser_defaite
            post_mortem_data = analyser_defaite(ao)
        except Exception as e:
            logger.warning(f"Erreur post-mortem pour {ao_id}: {e}")

    # Signaux faibles (marche fleche)
    signaux_faibles_data = None
    try:
        from signaux_faibles import detecter_signaux
        signaux_faibles_data = detecter_signaux(ao)
    except Exception as e:
        logger.warning(f"Erreur signaux faibles pour {ao_id}: {e}")

    # Groupement / co-traitance
    groupement_eval = None
    try:
        from groupement import evaluer_besoin_groupement
        groupement_eval = evaluer_besoin_groupement(ao)
    except Exception as e:
        logger.warning(f"Erreur groupement pour {ao_id}: {e}")

    # Analyse semantique DCE (si deja lancee)
    analyse_dce_complete = None
    try:
        clean_id_dce = ao_id.replace("BOAMP-", "").replace("PLACE-", "").replace("TED-", "").replace("MSEC-", "").replace("AWS-", "").replace("/", "_")
        fichier_analyse = DOSSIERS_GENERES_DIR / f"DCE_{clean_id_dce}" / "analyse_dce_complete.json"
        if not fichier_analyse.exists():
            fichier_analyse = DOSSIERS_GENERES_DIR / f"DCE_{ao_id.replace('/', '_')}" / "analyse_dce_complete.json"
        if fichier_analyse.exists():
            with open(fichier_analyse, "r", encoding="utf-8") as f:
                analyse_dce_complete = json.load(f)
    except Exception as e:
        logger.warning(f"Erreur chargement analyse DCE pour {ao_id}: {e}")

    # Prediction de victoire (scoring predictif)
    prediction_data = None
    try:
        from scoring_predictif import predire_victoire
        prediction_data = predire_victoire(ao)
    except Exception as e:
        logger.warning(f"Erreur scoring predictif pour {ao_id}: {e}")

    # Dossiers similaires (si pas de dossier genere)
    dossiers_similaires = []
    if not dossier_genere:
        try:
            from duplication_dossier import lister_dossiers_existants, trouver_dossier_similaire
            dossiers_existants = lister_dossiers_existants()
            if dossiers_existants:
                dossiers_similaires = trouver_dossier_similaire(ao, dossiers_existants)
        except Exception as e:
            logger.warning(f"Erreur dossiers similaires pour {ao_id}: {e}")

    return render_template("ao_detail.html", ao=ao, dossier=dossier_genere, note=note,
                           prestations=prestations, go_nogo=go_nogo,
                           type_presta=type_presta, modele=modele,
                           checklist=checklist, checklist_etat=checklist_etat,
                           commentaires=commentaires, estimation=estimation,
                           score_acheteur=score_acheteur_data,
                           ao_similaires=ao_similaires,
                           post_mortem=post_mortem_data,
                           signaux_faibles=signaux_faibles_data,
                           groupement=groupement_eval,
                           analyse_dce_complete=analyse_dce_complete,
                           prediction=prediction_data,
                           dossiers_similaires=dossiers_similaires,
                           decp=decp)


@app.route("/ao/<path:ao_id>/statut", methods=["POST"])
def changer_statut(ao_id):
    appels = charger_ao()
    nouveau_statut = request.form.get("statut", "nouveau")
    ao_trouve = None
    for ao in appels:
        if ao.get("id") == ao_id:
            ao["statut"] = nouveau_statut
            ao_trouve = ao
            break
    sauvegarder_ao(appels)
    # Memoire adaptative : indexer le memoire si AO gagne
    if nouveau_statut == "gagne" and ao_trouve:
        try:
            from memoire_adaptative import sauvegarder_memoire_gagnant
            sauvegarder_memoire_gagnant(ao_trouve)
        except Exception as e:
            logger.warning(f"Memoire adaptative: erreur sauvegarde pour {ao_id}: {e}")
    # Recalibrer le scoring predictif si gagne ou perdu
    if nouveau_statut in ("gagne", "perdu"):
        try:
            from scoring_predictif import calibrer_auto
            calibrer_auto()
        except Exception as e:
            logger.warning(f"Scoring predictif: erreur calibration pour {ao_id}: {e}")
    return redirect(url_for("detail_ao", ao_id=ao_id))


@app.route("/ao/<path:ao_id>/note", methods=["POST"])
def sauvegarder_note(ao_id):
    notes = charger_notes()
    note_text = request.form.get("note", "").strip()
    if note_text:
        notes[ao_id] = note_text
    elif ao_id in notes:
        del notes[ao_id]
    sauvegarder_notes(notes)
    return redirect(url_for("detail_ao", ao_id=ao_id))


@app.route("/ao/<path:ao_id>/checklist", methods=["POST"])
def maj_checklist(ao_id):
    """Met a jour l'etat d'un item de la checklist via AJAX."""
    data = request.get_json()
    if not data or "item_id" not in data:
        return jsonify({"error": "item_id requis"}), 400
    etat = _charger_checklist_etat(ao_id)
    etat[data["item_id"]] = data.get("checked", False)
    _sauvegarder_checklist_etat(ao_id, etat)
    return jsonify({"status": "ok"})


@app.route("/ao/<path:ao_id>/commentaire", methods=["POST"])
def ajouter_commentaire(ao_id):
    """Ajoute un commentaire sur un AO."""
    auteur = request.form.get("auteur", "").strip()
    texte = request.form.get("texte", "").strip()
    if not texte:
        return redirect(url_for("detail_ao", ao_id=ao_id))
    commentaire = {
        "auteur": auteur or "Anonyme",
        "texte": texte,
        "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
    }
    _sauvegarder_commentaire(ao_id, commentaire)
    return redirect(url_for("detail_ao", ao_id=ao_id))


@app.route("/ao/<path:ao_id>/telecharger-dce", methods=["POST"])
def telecharger_dce_route(ao_id):
    """POST - Tente de telecharger le DCE automatiquement."""
    appels = charger_ao()
    ao_dict = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao_dict:
        return jsonify({"error": "AO non trouve"}), 404

    def _dl_bg():
        try:
            from dce_auto import telecharger_dce_auto
            socketio.emit("veille_log", {"msg": f"Telechargement DCE: {ao_dict.get('titre', '')[:50]}..."})
            result = telecharger_dce_auto(ao_dict)
            if result["success"]:
                socketio.emit("veille_log", {"msg": f"DCE: {result['nb_fichiers']} fichier(s) telecharge(s)"})
            else:
                socketio.emit("veille_log", {"msg": f"DCE: {result['erreur']}"})
            socketio.emit("dce_complete", result)
        except Exception as e:
            socketio.emit("veille_log", {"msg": f"Erreur DCE: {e}"})
            socketio.emit("dce_complete", {"success": False, "erreur": str(e)})

    thread = threading.Thread(target=_dl_bg, daemon=True)
    thread.start()
    return jsonify({"status": "started"})


@app.route("/ao/<path:ao_id>/generer", methods=["POST"])
def generer_dossier(ao_id):
    """Lance la generation du dossier en arriere-plan via WebSocket."""
    appels = charger_ao()
    ao_dict = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao_dict:
        return jsonify({"error": "AO non trouve"}), 404

    def _generer_en_bg():
        try:
            # Essayer le generateur complet (local)
            try:
                from veille import AppelOffre
                from generateur import GenerateurMemoire
                from dce_downloader import telecharger_dce

                config = charger_config()
                ao = AppelOffre(**{k: v for k, v in ao_dict.items()
                                  if k in AppelOffre.__dataclass_fields__})

                socketio.emit("veille_log", {"msg": f"Generation dossier: {ao.titre[:60]}..."})

                dossier_ao = RESULTATS_DIR / f"AO_{ao.id}"
                dce_path = None
                try:
                    fichiers_dce = telecharger_dce(ao, dossier_ao, config=config)
                    if fichiers_dce:
                        dce_path = dossier_ao / "DCE"
                        socketio.emit("veille_log", {"msg": f"DCE telecharge: {len(fichiers_dce)} fichier(s)"})
                except Exception as e:
                    socketio.emit("veille_log", {"msg": f"DCE non disponible: {e}"})

                generateur = GenerateurMemoire(config)
                dossier = generateur.generer_dossier_complet(ao, dce_path=dce_path)
                socketio.emit("veille_log", {"msg": f"Dossier genere: {dossier.name}"})
                socketio.emit("generation_complete", {"ao_id": ao_id, "dossier": str(dossier)})
                return

            except ImportError:
                pass  # Pas en local, utiliser le generateur Render

            # Generateur Render (leger, memoire technique via API Claude)
            from generateur_render import generer_memoire_technique
            from modeles_reponse import detecter_type_prestation

            socketio.emit("veille_log", {"msg": f"Generation (Render): {ao_dict.get('titre', '')[:60]}..."})
            type_presta = detecter_type_prestation(ao_dict)
            socketio.emit("veille_log", {"msg": f"Type detecte: {type_presta}"})

            result = generer_memoire_technique(ao_dict, type_presta)
            if result["success"]:
                socketio.emit("veille_log", {"msg": f"Memoire genere: {result['nb_mots']} mots"})
                socketio.emit("generation_complete", {"ao_id": ao_id, "dossier": result["dossier_nom"]})
            else:
                socketio.emit("veille_log", {"msg": f"Erreur: {result['erreur']}"})
                socketio.emit("generation_complete", {"ao_id": ao_id, "error": result["erreur"]})

        except Exception as e:
            socketio.emit("veille_log", {"msg": f"Erreur generation: {e}"})
            socketio.emit("generation_complete", {"ao_id": ao_id, "error": str(e)})

    thread = threading.Thread(target=_generer_en_bg, daemon=True)
    thread.start()
    return jsonify({"status": "started", "ao_id": ao_id})


REVIEWS_FILE = DASHBOARD_DIR / "reviews.json"


def charger_reviews() -> dict:
    if not REVIEWS_FILE.exists():
        return {}
    with open(REVIEWS_FILE, "r", encoding="utf-8") as f:
        return json.load(f)


def sauvegarder_reviews(reviews: dict):
    REVIEWS_FILE.write_text(json.dumps(reviews, ensure_ascii=False, indent=2), encoding="utf-8")


@app.route("/dossiers")
def liste_dossiers():
    dossiers = lister_dossiers()
    reviews = charger_reviews()
    # Enrichir avec le statut de review
    for d in dossiers:
        d["review"] = reviews.get(d["nom"], {})
    return render_template("dossiers.html", dossiers=dossiers)


@app.route("/dossiers/<path:nom>")
def detail_dossier(nom):
    """Affiche les fichiers d'un dossier avec preview."""
    dossiers = lister_dossiers()
    dossier = next((d for d in dossiers if d["nom"] == nom), None)
    if not dossier:
        return redirect(url_for("liste_dossiers"))
    reviews = charger_reviews()
    review = reviews.get(nom, {"statut": "en_attente", "commentaires": []})
    return render_template("dossier_detail.html", dossier=dossier, review=review)


@app.route("/dossiers/<path:nom>/review", methods=["POST"])
def review_dossier(nom):
    """Ajoute un commentaire de relecture."""
    reviews = charger_reviews()
    if nom not in reviews:
        reviews[nom] = {"statut": "en_attente", "commentaires": []}
    data = request.form
    commentaire = data.get("commentaire", "").strip()
    auteur = data.get("auteur", "Anonyme").strip()
    statut = data.get("statut", "")
    if commentaire:
        reviews[nom]["commentaires"].append({
            "auteur": auteur,
            "texte": commentaire,
            "date": datetime.now().strftime("%Y-%m-%d %H:%M"),
        })
    if statut:
        reviews[nom]["statut"] = statut
    sauvegarder_reviews(reviews)
    return redirect(url_for("detail_dossier", nom=nom))


@app.route("/dossiers/<path:nom>/fichier/<path:fichier>")
def servir_fichier(nom, fichier):
    """Sert un fichier du dossier - rendu HTML pour .md/.json, download sinon."""
    from flask import send_file
    # Chercher dans les deux emplacements possibles
    fichier_path = None
    for base in [RESULTATS_DIR, DOSSIERS_GENERES_DIR]:
        dossier_path = base / nom
        if not dossier_path.exists():
            continue
        candidate = dossier_path / fichier
        if not candidate.exists() or not candidate.is_file():
            continue
        # Securite : verifier que le chemin reste dans le base dir
        try:
            candidate.resolve().relative_to(base.resolve())
        except ValueError:
            return "Acces refuse", 403
        fichier_path = candidate
        break

    if not fichier_path:
        return "Fichier non trouve", 404

    ext = fichier_path.suffix.lower()

    # Pour .md et .json : rendu dans un template HTML
    if ext in (".md", ".json"):
        contenu_brut = fichier_path.read_text(encoding="utf-8", errors="replace")
        if ext == ".md":
            contenu_html = _convertir_md_en_html(contenu_brut)
            mode = "markdown"
        else:
            try:
                data = json.loads(contenu_brut)
                contenu_html = json.dumps(data, indent=2, ensure_ascii=False)
            except Exception:
                contenu_html = contenu_brut
            mode = "json"
        return render_template(
            "fichier_view.html",
            nom_dossier=nom,
            nom_fichier=fichier,
            contenu_html=contenu_html,
            mode=mode,
        )

    # Pour .docx : convertir en HTML lisible
    if ext == ".docx":
        try:
            contenu_html = _convertir_docx_en_html(fichier_path)
            return render_template(
                "fichier_view.html",
                nom_dossier=nom,
                nom_fichier=fichier,
                contenu_html=contenu_html,
                mode="docx",
            )
        except Exception as e:
            logger.warning("Erreur conversion DOCX %s: %s", fichier, e)
            return send_file(str(fichier_path), as_attachment=True)

    # Pour .xlsx : convertir en tableau HTML
    if ext == ".xlsx":
        try:
            contenu_html = _convertir_xlsx_en_html(fichier_path)
            return render_template(
                "fichier_view.html",
                nom_dossier=nom,
                nom_fichier=fichier,
                contenu_html=contenu_html,
                mode="xlsx",
            )
        except Exception as e:
            logger.warning("Erreur conversion XLSX %s: %s", fichier, e)
            return send_file(str(fichier_path), as_attachment=True)

    # Pour les autres formats : servir directement
    return send_file(str(fichier_path), as_attachment=True)


def _convertir_docx_en_html(filepath) -> str:
    """Convertit un fichier DOCX en HTML basique pour affichage navigateur."""
    from docx import Document
    doc = Document(str(filepath))
    html_parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if "heading 1" in style or "titre 1" in style:
            html_parts.append(f"<h1>{text}</h1>")
        elif "heading 2" in style or "titre 2" in style:
            html_parts.append(f"<h2>{text}</h2>")
        elif "heading 3" in style or "titre 3" in style:
            html_parts.append(f"<h3>{text}</h3>")
        elif "heading" in style or "titre" in style:
            html_parts.append(f"<h4>{text}</h4>")
        else:
            # Gras / italique inline
            runs_html = []
            for run in para.runs:
                t = run.text
                if not t:
                    continue
                if run.bold and run.italic:
                    t = f"<strong><em>{t}</em></strong>"
                elif run.bold:
                    t = f"<strong>{t}</strong>"
                elif run.italic:
                    t = f"<em>{t}</em>"
                runs_html.append(t)
            html_parts.append(f"<p>{''.join(runs_html)}</p>")

    # Tables
    for table in doc.tables:
        html_parts.append("<table class='docx-table'>")
        for i, row in enumerate(table.rows):
            html_parts.append("<tr>")
            tag = "th" if i == 0 else "td"
            for cell in row.cells:
                html_parts.append(f"<{tag}>{cell.text.strip()}</{tag}>")
            html_parts.append("</tr>")
        html_parts.append("</table>")

    return "\n".join(html_parts)


def _convertir_xlsx_en_html(filepath) -> str:
    """Convertit un fichier XLSX en tableau HTML."""
    from openpyxl import load_workbook
    wb = load_workbook(str(filepath), read_only=True, data_only=True)
    html_parts = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        if len(wb.sheetnames) > 1:
            html_parts.append(f"<h3>{sheet_name}</h3>")
        html_parts.append("<table class='xlsx-table'>")
        for i, row in enumerate(ws.iter_rows(values_only=True)):
            html_parts.append("<tr>")
            tag = "th" if i == 0 else "td"
            for cell in row:
                val = str(cell) if cell is not None else ""
                # Formater les nombres
                if isinstance(cell, (int, float)):
                    try:
                        val = f"{cell:,.2f}".replace(",", " ").replace(".00", "")
                    except (ValueError, TypeError):
                        pass
                html_parts.append(f"<{tag}>{val}</{tag}>")
            html_parts.append("</tr>")
        html_parts.append("</table>")
    wb.close()
    return "\n".join(html_parts)


def _convertir_md_en_html(texte: str) -> str:
    """Conversion basique Markdown -> HTML (regex, sans librairie externe)."""
    import re

    # Proteger les blocs de code d'abord
    code_blocks = []
    def _save_code_block(m):
        code_blocks.append(m.group(1))
        return f"__CODE_BLOCK_{len(code_blocks) - 1}__"

    texte = re.sub(r"```(?:\w*)\n(.*?)```", _save_code_block, texte, flags=re.DOTALL)

    # Tableaux markdown
    lines = texte.split("\n")
    result_lines = []
    in_table = False
    for line in lines:
        stripped = line.strip()
        if stripped.startswith("|") and stripped.endswith("|"):
            cells = [c.strip() for c in stripped.strip("|").split("|")]
            # Ligne separateur (---|---|---) -> skip
            if all(set(c.strip()) <= {"-", ":"} for c in cells):
                continue
            if not in_table:
                result_lines.append("<table class='md-table'>")
                result_lines.append("<thead><tr>" + "".join(f"<th>{c}</th>" for c in cells) + "</tr></thead><tbody>")
                in_table = True
            else:
                result_lines.append("<tr>" + "".join(f"<td>{c}</td>" for c in cells) + "</tr>")
        else:
            if in_table:
                result_lines.append("</tbody></table>")
                in_table = False
            result_lines.append(line)
    if in_table:
        result_lines.append("</tbody></table>")
    texte = "\n".join(result_lines)

    # Headings (ordre decroissant)
    texte = re.sub(r"^######\s+(.+)$", r"<h6>\1</h6>", texte, flags=re.MULTILINE)
    texte = re.sub(r"^#####\s+(.+)$", r"<h5>\1</h5>", texte, flags=re.MULTILINE)
    texte = re.sub(r"^####\s+(.+)$", r"<h4>\1</h4>", texte, flags=re.MULTILINE)
    texte = re.sub(r"^###\s+(.+)$", r"<h3>\1</h3>", texte, flags=re.MULTILINE)
    texte = re.sub(r"^##\s+(.+)$", r"<h2>\1</h2>", texte, flags=re.MULTILINE)
    texte = re.sub(r"^#\s+(.+)$", r"<h1>\1</h1>", texte, flags=re.MULTILINE)

    # Horizontal rules
    texte = re.sub(r"^---+$", "<hr>", texte, flags=re.MULTILINE)

    # Bold et italic
    texte = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", texte)
    texte = re.sub(r"\*(.+?)\*", r"<em>\1</em>", texte)

    # Listes a puces
    lines = texte.split("\n")
    result_lines = []
    in_list = False
    for line in lines:
        stripped = line.strip()
        if re.match(r"^[-*]\s+", stripped):
            if not in_list:
                result_lines.append("<ul>")
                in_list = True
            item = re.sub(r"^[-*]\s+", "", stripped)
            result_lines.append(f"<li>{item}</li>")
        else:
            if in_list:
                result_lines.append("</ul>")
                in_list = False
            result_lines.append(line)
    if in_list:
        result_lines.append("</ul>")
    texte = "\n".join(result_lines)

    # Listes numerotees
    lines = texte.split("\n")
    result_lines = []
    in_ol = False
    for line in lines:
        stripped = line.strip()
        if re.match(r"^\d+\.\s+", stripped):
            if not in_ol:
                result_lines.append("<ol>")
                in_ol = True
            item = re.sub(r"^\d+\.\s+", "", stripped)
            result_lines.append(f"<li>{item}</li>")
        else:
            if in_ol:
                result_lines.append("</ol>")
                in_ol = False
            result_lines.append(line)
    if in_ol:
        result_lines.append("</ol>")
    texte = "\n".join(result_lines)

    # Restaurer les blocs de code
    for i, block in enumerate(code_blocks):
        escaped = block.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        texte = texte.replace(f"__CODE_BLOCK_{i}__", f"<pre><code>{escaped}</code></pre>")

    # Paragraphes : lignes non-vides non-HTML -> <p>
    lines = texte.split("\n")
    result_lines = []
    for line in lines:
        stripped = line.strip()
        if stripped and not stripped.startswith("<"):
            result_lines.append(f"<p>{stripped}</p>")
        elif stripped == "":
            result_lines.append("")
        else:
            result_lines.append(line)

    return "\n".join(result_lines)


@app.route("/dossiers/<path:nom>/fichier/<path:fichier>/raw")
def servir_fichier_raw(nom, fichier):
    """Telecharge le fichier brut (sans rendu HTML)."""
    from flask import send_file
    for base in [RESULTATS_DIR, DOSSIERS_GENERES_DIR]:
        dossier_path = base / nom
        if not dossier_path.exists():
            continue
        candidate = dossier_path / fichier
        if not candidate.exists() or not candidate.is_file():
            continue
        try:
            candidate.resolve().relative_to(base.resolve())
        except ValueError:
            return "Acces refuse", 403
        return send_file(str(candidate), as_attachment=True)
    return "Fichier non trouve", 404


@app.route("/dossiers/<path:nom>/zip")
def export_dossier_zip(nom):
    """Telecharge tout le dossier en ZIP."""
    import zipfile
    from flask import send_file

    # Trouver le dossier
    dossier_path = None
    for base in [RESULTATS_DIR, DOSSIERS_GENERES_DIR]:
        candidate = base / nom
        if candidate.exists() and candidate.is_dir():
            dossier_path = candidate
            break

    if not dossier_path:
        return "Dossier non trouve", 404

    # Creer le ZIP en memoire
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zf:
        for fichier in dossier_path.rglob("*"):
            if fichier.is_file():
                arcname = fichier.relative_to(dossier_path)
                zf.write(fichier, arcname)

    zip_buffer.seek(0)
    return send_file(
        zip_buffer,
        mimetype="application/zip",
        as_attachment=True,
        download_name=f"{nom}.zip",
    )


# --- Export PDF (vue impression navigateur) ---

ORDRE_FICHIERS = [
    ("analyse", "Analyse Go/No-Go"),
    ("memoire_technique", "Memoire Technique"),
    ("lettre_candidature", "Lettre de Candidature"),
    ("programme", "Programme de Formation"),
    ("bpu", "Bordereau de Prix"),
    ("dpgf", "Bordereau de Prix"),
    ("planning", "Planning Previsionnel"),
    ("cv", "CV des Formateurs"),
    ("references", "References Clients"),
    ("dc1", "Formulaires DC1/DC2"),
    ("dc2", "Formulaires DC1/DC2"),
    ("moyens", "Moyens Techniques"),
    ("acte_engagement", "Acte d'Engagement"),
    ("dume", "DUME"),
    ("checklist", "Checklist de Soumission"),
]


def _ordre_fichier(nom_fichier: str) -> int:
    """Retourne un index de tri pour ordonner les fichiers du dossier."""
    nom_lower = nom_fichier.lower()
    for i, (prefix, _) in enumerate(ORDRE_FICHIERS):
        if prefix in nom_lower:
            return i
    return 999


def _titre_fichier(nom_fichier: str) -> str:
    """Retourne un titre lisible a partir du nom de fichier."""
    nom_lower = nom_fichier.lower()
    for prefix, titre in ORDRE_FICHIERS:
        if prefix in nom_lower:
            return titre
    # Fallback : nettoyer le nom
    base = Path(nom_fichier).stem
    return base.replace("_", " ").replace("-", " ").title()


@app.route("/dossiers/<path:nom>/print")
def dossier_print_view(nom):
    """Vue d'impression complete du dossier - tous les fichiers concatenes."""
    # Trouver le dossier
    dossier_path = None
    for base in [RESULTATS_DIR, DOSSIERS_GENERES_DIR]:
        candidate = base / nom
        if candidate.exists() and candidate.is_dir():
            dossier_path = candidate
            break

    if not dossier_path:
        return "Dossier non trouve", 404

    # Collecter et trier les fichiers
    fichiers = sorted(
        [f for f in dossier_path.iterdir() if f.is_file()],
        key=lambda f: _ordre_fichier(f.name),
    )

    # Convertir chaque fichier en HTML
    documents = []
    for f in fichiers:
        ext = f.suffix.lower()
        titre = _titre_fichier(f.name)
        contenu_html = None

        if ext == ".md":
            contenu_brut = f.read_text(encoding="utf-8", errors="replace")
            contenu_html = _convertir_md_en_html(contenu_brut)
        elif ext == ".docx":
            try:
                contenu_html = _convertir_docx_en_html(f)
            except Exception as e:
                logger.warning("Erreur conversion DOCX pour print %s: %s", f.name, e)
                contenu_html = f"<p><em>Fichier DOCX non convertible : {f.name}</em></p>"
        elif ext == ".xlsx":
            try:
                contenu_html = _convertir_xlsx_en_html(f)
            except Exception as e:
                logger.warning("Erreur conversion XLSX pour print %s: %s", f.name, e)
                contenu_html = f"<p><em>Fichier Excel non convertible : {f.name}</em></p>"
        elif ext == ".json":
            try:
                data = json.loads(f.read_text(encoding="utf-8", errors="replace"))
                contenu_html = f"<pre>{json.dumps(data, indent=2, ensure_ascii=False)}</pre>"
            except Exception:
                pass
        elif ext == ".txt":
            contenu_brut = f.read_text(encoding="utf-8", errors="replace")
            contenu_html = f"<pre>{contenu_brut}</pre>"

        if contenu_html:
            documents.append({"titre": titre, "nom_fichier": f.name, "contenu_html": contenu_html})

    # Extraire le titre de l'AO depuis le nom du dossier
    titre_ao = nom.replace("AO_", "").replace("_", " ")

    return render_template(
        "dossier_print.html",
        nom=nom,
        titre_ao=titre_ao,
        documents=documents,
        now=datetime.now().strftime("%d/%m/%Y"),
    )


@app.route("/export/csv")
def export_csv():
    """Exporte les AO filtres en CSV."""
    appels = charger_ao()
    source = request.args.get("source", "")
    score_min_str = request.args.get("score_min", "0")
    statut = request.args.get("statut", "")
    recherche = request.args.get("q", "")
    tri = request.args.get("tri", "score")

    try:
        score_min = float(score_min_str)
    except ValueError:
        score_min = 0

    filtre = filtrer_ao(appels, source, score_min, statut, recherche, tri)

    output = io.StringIO()
    writer = csv.writer(output, delimiter=";")
    writer.writerow(["ID", "Source", "Titre", "Acheteur", "Score", "Date limite",
                     "Region", "Budget", "Statut", "URL", "Description"])
    for ao in filtre:
        writer.writerow([
            ao.get("id", ""),
            ao.get("source", ""),
            ao.get("titre", ""),
            ao.get("acheteur", ""),
            f"{(ao.get('score_pertinence') or 0) * 100:.0f}%",
            ao.get("date_limite", ""),
            ao.get("region", ""),
            ao.get("budget_estime", ""),
            ao.get("statut", ""),
            ao.get("url", ""),
            (ao.get("description") or "")[:200],
        ])

    resp = make_response(output.getvalue())
    resp.headers["Content-Type"] = "text/csv; charset=utf-8"
    resp.headers["Content-Disposition"] = f"attachment; filename=ao_hunter_export_{datetime.now().strftime('%Y%m%d')}.csv"
    return resp


# --- API REST ---

@app.route("/api/ao")
def api_ao_list():
    """GET /api/ao - Liste tous les AO avec filtres optionnels."""
    appels = charger_ao()
    source = request.args.get("source", "")
    score_min = float(request.args.get("score_min", "0"))
    statut = request.args.get("statut", "")
    q = request.args.get("q", "")
    limit = int(request.args.get("limit", "50"))
    offset = int(request.args.get("offset", "0"))

    filtre = filtrer_ao(appels, source, score_min, statut, q)
    return jsonify({
        "total": len(filtre),
        "offset": offset,
        "limit": limit,
        "data": filtre[offset:offset + limit],
    })


@app.route("/api/ao/<path:ao_id>")
def api_ao_detail(ao_id):
    """GET /api/ao/<id> - Detail d'un AO."""
    appels = charger_ao()
    ao = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao:
        return jsonify({"error": "AO non trouve"}), 404
    notes = charger_notes()
    ao["note"] = notes.get(ao_id, "")
    return jsonify(ao)


@app.route("/api/ao/<path:ao_id>", methods=["PATCH"])
def api_ao_update(ao_id):
    """PATCH /api/ao/<id> - Met a jour statut et/ou note."""
    appels = charger_ao()
    data = request.get_json()
    if not data:
        return jsonify({"error": "JSON body requis"}), 400

    for ao in appels:
        if ao.get("id") == ao_id:
            if "statut" in data:
                ao["statut"] = data["statut"]
            sauvegarder_ao(appels)

            if "note" in data:
                notes = charger_notes()
                if data["note"]:
                    notes[ao_id] = data["note"]
                elif ao_id in notes:
                    del notes[ao_id]
                sauvegarder_notes(notes)

            return jsonify({"status": "ok", "id": ao_id})

    return jsonify({"error": "AO non trouve"}), 404


@app.route("/api/ao/<path:ao_id>/generer", methods=["POST"])
def api_generer(ao_id):
    """POST /api/ao/<id>/generer - Lance la generation du dossier."""
    return generer_dossier(ao_id)


@app.route("/api/ao/<path:ao_id>/dossiers-similaires")
def api_dossiers_similaires(ao_id):
    """GET /api/ao/<id>/dossiers-similaires - Dossiers existants similaires."""
    appels = charger_ao()
    ao = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao:
        return jsonify({"error": "AO non trouve"}), 404
    from duplication_dossier import lister_dossiers_existants, trouver_dossier_similaire
    dossiers = lister_dossiers_existants()
    similaires = trouver_dossier_similaire(ao, dossiers)
    return jsonify(similaires)


@app.route("/api/ao/<path:ao_id>/dupliquer-dossier", methods=["POST"])
def api_dupliquer_dossier(ao_id):
    """POST /api/ao/<id>/dupliquer-dossier - Duplique un dossier depuis un AO source."""
    appels = charger_ao()
    ao_cible = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao_cible:
        return jsonify({"error": "AO cible non trouve"}), 404

    data = request.get_json()
    if not data or "source_ao_id" not in data:
        return jsonify({"error": "source_ao_id requis"}), 400

    source_ao_id = data["source_ao_id"]

    from duplication_dossier import lister_dossiers_existants, dupliquer_dossier as _dupliquer
    dossiers = lister_dossiers_existants()
    dossier_source = next((d for d in dossiers if d["ao_id"] == source_ao_id), None)
    if not dossier_source:
        return jsonify({"error": f"Aucun dossier trouve pour l'AO source {source_ao_id}"}), 404

    resultat = _dupliquer(source_ao_id, ao_cible, dossier_source["dossier_path"])
    if "error" in resultat:
        return jsonify(resultat), 500
    return jsonify(resultat)


@app.route("/api/ao/<path:ao_id>/analyser-dce", methods=["POST"])
def api_analyser_dce(ao_id):
    """POST /api/ao/<id>/analyser-dce - Analyse semantique complete du DCE."""
    appels = charger_ao()
    ao_dict = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao_dict:
        return jsonify({"error": "AO non trouve"}), 404

    # Trouver le dossier DCE
    clean_id = ao_id.replace("BOAMP-", "").replace("PLACE-", "").replace("TED-", "").replace("MSEC-", "").replace("AWS-", "").replace("/", "_")
    dossier_dce = DOSSIERS_GENERES_DIR / f"DCE_{clean_id}"

    if not dossier_dce.exists():
        # Essayer aussi avec l'ID brut
        dossier_dce = DOSSIERS_GENERES_DIR / f"DCE_{ao_id.replace('/', '_')}"
    if not dossier_dce.exists():
        return jsonify({"error": "DCE non telecharge. Telechargez d'abord le DCE."}), 400

    def _analyser_bg():
        try:
            socketio.emit("veille_log", {"msg": "Analyse semantique du DCE en cours..."})
            from analyse_semantique_dce import analyser_dce_complet_ia
            resultat = analyser_dce_complet_ia(dossier_dce, ao_dict)
            if "erreur" in resultat:
                socketio.emit("veille_log", {"msg": f"Erreur analyse DCE: {resultat['erreur']}"})
                socketio.emit("analyse_dce_complete", {"success": False, "erreur": resultat["erreur"]})
            else:
                score = resultat.get("score_adequation", 0)
                socketio.emit("veille_log", {"msg": f"Analyse DCE terminee - Score: {score}/100"})
                socketio.emit("analyse_dce_complete", {"success": True, "score": score})
        except Exception as e:
            logger.error(f"Erreur analyse semantique DCE: {e}")
            socketio.emit("veille_log", {"msg": f"Erreur analyse DCE: {e}"})
            socketio.emit("analyse_dce_complete", {"success": False, "erreur": str(e)})

    thread = threading.Thread(target=_analyser_bg, daemon=True)
    thread.start()
    return jsonify({"status": "started"})


@app.route("/api/stats")
def api_stats():
    """GET /api/stats - Statistiques globales."""
    appels = charger_ao()
    return jsonify(stats_ao(appels))


@app.route("/api/dossiers")
def api_dossiers():
    """GET /api/dossiers - Liste des dossiers generes."""
    return jsonify(lister_dossiers())


@app.route("/api/urgents")
def api_urgents():
    """GET /api/urgents - AO avec deadline proche (pour notifications)."""
    appels = charger_ao()
    return jsonify(_ao_urgents(appels, max_days=3))


@app.route("/api/concurrence")
def api_concurrence():
    """GET /api/concurrence - Top concurrents et attributions recentes."""
    from veille_concurrence import charger_concurrence, analyser_concurrents
    attributions = charger_concurrence()
    concurrents = analyser_concurrents(attributions)
    return jsonify({
        "total_attributions": len(attributions),
        "top_concurrents": concurrents[:15],
        "dernieres_attributions": attributions[-10:][::-1],
    })


@app.route("/api/historique-veille")
def api_historique_veille():
    """GET /api/historique-veille - Historique des cycles de veille + tendances."""
    from veille_render import charger_historique, stats_tendances
    historique = charger_historique()
    tendances = stats_tendances(historique)
    return jsonify(tendances)


@app.route("/api/deadlines")
def api_deadlines():
    """GET /api/deadlines - AO avec deadlines a surveiller (J-7)."""
    from rappels import verifier_deadlines
    appels = charger_ao()
    alertes = verifier_deadlines(appels)
    return jsonify([{
        "id": a["ao"].get("id"),
        "titre": a["ao"].get("titre"),
        "acheteur": a["ao"].get("acheteur"),
        "statut": a["ao"].get("statut"),
        "date_limite": a["date_limite"],
        "jours_restants": a["jours_restants"],
        "urgence": a["urgence"],
    } for a in alertes])


@app.route("/api/pipeline")
def api_pipeline():
    """GET /api/pipeline - Statut et historique du pipeline automatique."""
    try:
        from pipeline_auto import charger_log
        log = charger_log()
        aujourd_hui = datetime.now().strftime("%Y-%m-%d")
        generes_today = sum(1 for e in log if e.get("action") == "GENERE" and e.get("timestamp", "").startswith(aujourd_hui))
        total_generes = sum(1 for e in log if e.get("action") == "GENERE")
        total_skip = sum(1 for e in log if e.get("action") == "SKIP")
        return jsonify({
            "generes_aujourd_hui": generes_today,
            "max_par_jour": 1,
            "total_generes": total_generes,
            "total_skip": total_skip,
            "derniers": log[-10:][::-1],
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/estimation/<path:ao_id>")
def api_estimation(ao_id):
    """GET /api/estimation/<id> - Estimation budget/concurrence/accessibilite."""
    appels = charger_ao()
    ao = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao:
        return jsonify({"error": "AO non trouve"}), 404
    try:
        from estimation_marche import estimer_marche
        return jsonify(estimer_marche(ao))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/veille", methods=["POST"])
def api_lancer_veille():
    """POST /api/veille - Lance une veille en arriere-plan.
    Sur Render: utilise veille_render (BOAMP+TED lightweight).
    En local: utilise le module complet veille+filtre."""
    def _veille_bg():
        try:
            socketio.emit("veille_log", {"msg": "Recherche en cours..."})

            # Sur Render ou si modules complets non dispo: veille legere
            try:
                config = charger_config()
                from veille import Veilleur
                from filtre import FiltreAO
                # Mode local complet
                veilleur = Veilleur(config)
                tous_ao = veilleur.lancer_recherche()
                socketio.emit("veille_log", {"msg": f"{len(tous_ao)} AO bruts trouves"})

                filtre = FiltreAO(config)
                pertinents = filtre.filtrer(tous_ao)
                socketio.emit("veille_log", {"msg": f"{len(pertinents)} AO pertinents"})

                existants = charger_ao()
                ids_existants = {ao["id"] for ao in existants}
                nouveaux = 0
                for ao in pertinents:
                    d = ao.to_dict()
                    if d["id"] not in ids_existants:
                        existants.append(d)
                        nouveaux += 1
                sauvegarder_ao(existants)
            except ImportError:
                # Mode Render: veille legere
                from veille_render import lancer_veille
                result = lancer_veille()
                nouveaux = result["nouveaux"]
                existants = charger_ao()

            socketio.emit("veille_log", {"msg": f"Termine: {nouveaux} nouveaux AO ajoutes"})
            socketio.emit("veille_complete", {"nouveaux": nouveaux, "total": len(existants)})
        except Exception as e:
            socketio.emit("veille_log", {"msg": f"Erreur: {e}"})
            socketio.emit("veille_complete", {"error": str(e)})

    thread = threading.Thread(target=_veille_bg, daemon=True)
    thread.start()
    return jsonify({"status": "started"})


@app.route("/api/roi")
def api_roi():
    """GET /api/roi - Stats ROI completes en JSON."""
    stats = _calculer_roi_stats()
    # Convertir top_acheteurs (list of tuples) en list of dicts pour JSON
    stats["top_acheteurs"] = [{"acheteur": nom, **data} for nom, data in stats["top_acheteurs"]]
    return jsonify(stats)


@app.route("/api/veille/status")
def api_veille_status():
    """GET /api/veille/status - Statut du scheduler et derniere veille."""
    status = {"scheduler_actif": _scheduler is not None and _scheduler.running}
    if _scheduler and _scheduler.running:
        job = _scheduler.get_job("veille_auto")
        if job:
            status["prochaine_execution"] = str(job.next_run_time)
    return jsonify(status)


@app.route("/api/attributions")
def api_attributions():
    """GET /api/attributions - Historique des resultats d'attribution."""
    from veille_attributions import charger_attributions
    attributions = charger_attributions()
    return jsonify({
        "total": len(attributions),
        "attributions": attributions,
        "gagnes": [a for a in attributions if a.get("notre_statut") == "gagne"],
        "perdus": [a for a in attributions if a.get("notre_statut") == "perdu"],
    })


@app.route("/api/acheteurs-cles")
def api_acheteurs_cles():
    """GET /api/acheteurs-cles - Top acheteurs les plus compatibles avec Almera."""
    from score_acheteur import top_acheteurs
    n = int(request.args.get("n", "15"))
    return jsonify(top_acheteurs(n=n))


# --- API Prix ---

@app.route("/api/prix/stats")
def api_prix_stats():
    """GET /api/prix/stats - Statistiques de prix par type de prestation."""
    from modeles_prix import stats_prix
    type_prestation = request.args.get("type")
    return jsonify(stats_prix(type_prestation=type_prestation))


@app.route("/api/prix/enregistrer", methods=["POST"])
def api_prix_enregistrer():
    """POST /api/prix/enregistrer - Enregistrer un prix soumis."""
    from modeles_prix import enregistrer_prix
    data = request.get_json(force=True)
    ao_id = data.get("ao_id")
    type_prestation = data.get("type_prestation", "formation_intra")
    montant = data.get("montant")

    if not ao_id or not montant:
        return jsonify({"error": "ao_id et montant requis"}), 400

    try:
        montant = float(montant)
    except (ValueError, TypeError):
        return jsonify({"error": "montant doit etre un nombre"}), 400

    entree = enregistrer_prix(
        ao_id=ao_id,
        type_prestation=type_prestation,
        montant=montant,
        nb_jours=data.get("nb_jours"),
        nb_personnes=data.get("nb_personnes"),
        resultat=data.get("resultat"),
    )
    return jsonify({"status": "ok", "entree": entree})


# --- API Post-mortem ---

@app.route("/api/post-mortem/stats")
def api_post_mortem_stats():
    """GET /api/post-mortem/stats - Statistiques agregees des AO perdus."""
    from post_mortem import stats_post_mortem
    return jsonify(stats_post_mortem())


# --- API Benchmark Prix ---

@app.route("/api/benchmark")
def api_benchmark():
    """GET /api/benchmark - Stats benchmark prix du marche."""
    from benchmark_prix import analyser_benchmark
    type_prestation = request.args.get("type", "formation")
    return jsonify(analyser_benchmark(type_prestation=type_prestation))


@app.route("/api/benchmark/position")
def api_benchmark_position():
    """GET /api/benchmark/position?prix=15000&type=formation - Positionnement prix."""
    from benchmark_prix import positionner_prix
    prix = request.args.get("prix")
    type_prestation = request.args.get("type", "formation")
    if not prix:
        return jsonify({"error": "parametre 'prix' requis"}), 400
    try:
        prix = float(prix)
    except (ValueError, TypeError):
        return jsonify({"error": "prix doit etre un nombre"}), 400
    return jsonify(positionner_prix(notre_prix=prix, type_prestation=type_prestation))


# --- API Soumission ---

@app.route("/api/ao/<path:ao_id>/preparer-soumission", methods=["POST"])
def api_preparer_soumission(ao_id):
    """POST /api/ao/<id>/preparer-soumission - Prepare le dossier pour soumission."""
    appels = charger_ao()
    ao = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao:
        return jsonify({"error": "AO non trouve"}), 404

    try:
        from soumission_helper import preparer_soumission
        result = preparer_soumission(ao)
        return jsonify(result)
    except Exception as e:
        logger.error(f"Erreur preparation soumission: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/ao/<path:ao_id>/guide-soumission")
def api_guide_soumission(ao_id):
    """GET /api/ao/<id>/guide-soumission - Guide de soumission pour la plateforme."""
    appels = charger_ao()
    ao = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao:
        return jsonify({"error": "AO non trouve"}), 404

    try:
        from soumission_helper import identifier_plateforme, generer_guide_soumission
        plateforme = identifier_plateforme(ao)
        guide = generer_guide_soumission(plateforme["id"])
        return jsonify({
            "plateforme": plateforme["nom"],
            "plateforme_id": plateforme["id"],
            "url_depot": plateforme["url_depot"],
            "etapes": guide,
        })
    except Exception as e:
        logger.error(f"Erreur guide soumission: {e}")
        return jsonify({"error": str(e)}), 500


# --- API Questions acheteur ---

@app.route("/api/ao/<path:ao_id>/questions", methods=["POST"])
def api_generer_questions(ao_id):
    """POST /api/ao/<id>/questions - Genere des questions a poser a l'acheteur."""
    appels = charger_ao()
    ao = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao:
        return jsonify({"error": "AO non trouve"}), 404

    try:
        from questions_acheteur import generer_questions, formater_questions_email
        result = generer_questions(ao)

        # Ajouter l'email formate si des questions ont ete generees
        if result.get("questions"):
            result["email_formate"] = formater_questions_email(result["questions"], ao)

        return jsonify(result)
    except Exception as e:
        logger.error(f"Erreur generation questions: {e}")
        return jsonify({"error": str(e)}), 500


# --- Pre-informations ---

@app.route("/preinfos")
def preinfos_view():
    """Page listant les pre-informations detectees."""
    from veille_preinformation import preinfos_actives, recommander_actions
    preinfos = preinfos_actives()
    for p in preinfos:
        p["actions"] = recommander_actions(p)
    return render_template("preinfos.html", preinfos=preinfos)


@app.route("/api/preinfos")
def api_preinfos():
    """GET /api/preinfos - JSON des pre-infos actives."""
    from veille_preinformation import preinfos_actives, recommander_actions
    preinfos = preinfos_actives()
    for p in preinfos:
        p["actions"] = recommander_actions(p)
    return jsonify({"total": len(preinfos), "preinfos": preinfos})


# --- CRM Acheteurs ---

@app.route("/crm")
def crm_view():
    """Page CRM avec la liste des acheteurs."""
    from crm_acheteurs import lister_tous_acheteurs, acheteurs_a_relancer
    acheteurs = lister_tous_acheteurs()
    relancer = acheteurs_a_relancer()
    return render_template("crm.html", acheteurs=acheteurs, a_relancer=relancer)


@app.route("/crm/<path:nom>")
def crm_fiche_view(nom):
    """Fiche detaillee d'un acheteur."""
    from crm_acheteurs import get_fiche
    fiche = get_fiche(nom, is_cle=True)
    if not fiche:
        return redirect(url_for("crm_view"))
    return render_template("crm_fiche.html", fiche=fiche, cle=nom)


@app.route("/crm/<path:nom>/note", methods=["POST"])
def crm_ajouter_note(nom):
    """Ajouter une note a un acheteur."""
    from crm_acheteurs import ajouter_note
    note = request.form.get("note", "").strip()
    if note:
        ajouter_note(nom, note)
    return redirect(url_for("crm_fiche_view", nom=nom))


@app.route("/crm/<path:nom>/contact", methods=["POST"])
def crm_ajouter_contact(nom):
    """Ajouter un contact a un acheteur."""
    from crm_acheteurs import ajouter_contact
    contact_info = {
        "nom": request.form.get("nom", "").strip(),
        "email": request.form.get("email", "").strip(),
        "telephone": request.form.get("telephone", "").strip(),
        "fonction": request.form.get("fonction", "").strip(),
    }
    if contact_info["nom"] or contact_info["email"]:
        ajouter_contact(nom, contact_info)
    return redirect(url_for("crm_fiche_view", nom=nom))


@app.route("/api/crm")
def api_crm():
    """GET /api/crm - JSON de tous les acheteurs."""
    from crm_acheteurs import lister_tous_acheteurs, acheteurs_a_relancer, top_acheteurs_actifs
    return jsonify({
        "acheteurs": lister_tous_acheteurs(),
        "a_relancer": acheteurs_a_relancer(),
        "top_actifs": top_acheteurs_actifs(20),
    })


# --- API Groupement ---

@app.route("/api/ao/<path:ao_id>/groupement", methods=["POST"])
def api_groupement(ao_id):
    """POST /api/ao/<id>/groupement - Genere les documents de groupement."""
    appels = charger_ao()
    ao = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao:
        return jsonify({"error": "AO non trouve"}), 404

    try:
        from groupement import generer_documents_groupement, suggestions_partenaires
        data = request.get_json(force=True) if request.is_json else {}
        partenaire_info = data.get("partenaire") if data else None
        documents = generer_documents_groupement(ao, partenaire_info=partenaire_info)
        suggestions = suggestions_partenaires(ao)
        return jsonify({
            "status": "ok",
            "documents": documents,
            "suggestions_partenaires": suggestions,
        })
    except Exception as e:
        logger.error(f"Erreur generation groupement: {e}")
        return jsonify({"error": str(e)}), 500


# --- Objectifs CA ---

@app.route("/objectifs")
def objectifs_view():
    """Page objectifs CA avec progression et recommandations."""
    from objectifs_ca import progression as calc_progression, recommander_pipeline, alertes_objectif
    from datetime import date as _date

    prog = calc_progression()
    recommandation = recommander_pipeline()
    alertes = alertes_objectif()

    return render_template("objectifs.html",
                           prog=prog,
                           objectifs=prog["objectifs"],
                           recommandation=recommandation,
                           alertes=alertes,
                           stats_par_mois=prog["stats_par_mois"],
                           mois_actuel=_date.today().month)


@app.route("/api/objectifs")
def api_objectifs_get():
    """GET /api/objectifs - Donnees objectifs CA."""
    from objectifs_ca import progression as calc_progression, recommander_pipeline, alertes_objectif
    prog = calc_progression()
    recommandation = recommander_pipeline()
    alertes = alertes_objectif()
    return jsonify({
        "progression": prog,
        "recommandation": recommandation,
        "alertes": alertes,
    })


@app.route("/api/objectifs", methods=["POST"])
def api_objectifs_post():
    """POST /api/objectifs - Definir l'objectif annuel."""
    from objectifs_ca import definir_objectif
    data = request.get_json(force=True)
    annuel = data.get("objectif_annuel")
    if not annuel:
        return jsonify({"error": "objectif_annuel requis"}), 400
    try:
        annuel = float(annuel)
    except (ValueError, TypeError):
        return jsonify({"error": "objectif_annuel doit etre un nombre"}), 400
    marge = data.get("marge_cible_pct", 30)
    result = definir_objectif(annuel, marge_cible_pct=float(marge))
    return jsonify({"status": "ok", "objectifs": result})


@app.route("/api/memoires-gagnants")
def api_memoires_gagnants():
    """Liste des memoires techniques indexes (modeles gagnants)."""
    try:
        from memoire_adaptative import _charger_index
        index = _charger_index()
        return jsonify(index)
    except Exception as e:
        logger.error(f"Erreur chargement memoires gagnants: {e}")
        return jsonify([])


# --- Scoring predictif API ---

@app.route("/api/scoring/stats")
def api_scoring_stats():
    """GET /api/scoring/stats - Stats du modele de scoring predictif."""
    try:
        from scoring_predictif import stats_modele
        return jsonify(stats_modele())
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/scoring/prediction/<path:ao_id>")
def api_scoring_prediction(ao_id):
    """GET /api/scoring/prediction/<ao_id> - Prediction pour un AO."""
    appels = charger_ao()
    ao = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao:
        return jsonify({"error": "AO non trouve"}), 404
    try:
        from scoring_predictif import predire_victoire
        return jsonify(predire_victoire(ao))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- WebSocket ---

@socketio.on("connect")
def on_connect():
    logger.info("Client WebSocket connecte")


# --- Rapport hebdomadaire ---

@app.route("/api/rapport-hebdo", methods=["POST"])
def api_rapport_hebdo_forcer():
    """POST /api/rapport-hebdo - Forcer la generation manuelle du rapport hebdomadaire."""
    try:
        from rapport_hebdo import rapport_et_envoi
        result = rapport_et_envoi()
        return jsonify(result)
    except Exception as e:
        return jsonify({"erreur": str(e)}), 500


@app.route("/api/rapport-hebdo/dernier")
def api_rapport_hebdo_dernier():
    """GET /api/rapport-hebdo/dernier - Retourner le dernier rapport hebdomadaire."""
    try:
        from rapport_hebdo import _charger_rapports
        rapports = _charger_rapports()
        if not rapports:
            return jsonify({"erreur": "Aucun rapport disponible"}), 404
        return jsonify(rapports[-1])
    except Exception as e:
        return jsonify({"erreur": str(e)}), 500


# --- Alertes instantanees ---

@app.route("/api/alertes")
def api_alertes():
    """GET /api/alertes - Alertes instantanees recentes (AO a fort potentiel)."""
    try:
        from alertes_instantanees import get_alertes_recentes
        nb = request.args.get("nb", 20, type=int)
        alertes = get_alertes_recentes(nb)
        return jsonify({"alertes": alertes, "total": len(alertes)})
    except Exception as e:
        return jsonify({"erreur": str(e)}), 500


# --- Sync endpoints ---

_app_start_time = datetime.now()


@app.route("/api/sync/status")
def api_sync_status():
    """GET /api/sync/status - Etat de synchronisation et sante du dashboard."""
    appels = charger_ao()
    uptime_seconds = (datetime.now() - _app_start_time).total_seconds()

    # Derniere synchro (depuis le fichier meta si present)
    derniere_synchro = None
    meta_file = DASHBOARD_DIR / ".sync_meta.json"
    if meta_file.exists():
        try:
            meta = json.loads(meta_file.read_text(encoding="utf-8"))
            derniere_synchro = meta.get("dernier_pull")
        except Exception:
            pass

    return jsonify({
        "derniere_synchro": derniere_synchro,
        "nb_ao": len(appels),
        "version": "AO Hunter Dashboard v2",
        "uptime": round(uptime_seconds),
        "uptime_humain": f"{int(uptime_seconds // 3600)}h {int((uptime_seconds % 3600) // 60)}m",
        "environnement": "render" if os.environ.get("RENDER") else "local",
    })


@app.route("/api/sync/pull", methods=["POST"])
def api_sync_pull():
    """POST /api/sync/pull - Force un rechargement des donnees depuis les fichiers."""
    try:
        # Recharger les fichiers de donnees (utile apres un push externe)
        appels = charger_ao()
        concurrence_file = DASHBOARD_DIR / "concurrence.json"
        historique_file = DASHBOARD_DIR / "historique_veille.json"

        nb_concurrence = 0
        if concurrence_file.exists():
            try:
                data = json.loads(concurrence_file.read_text(encoding="utf-8"))
                nb_concurrence = len(data) if isinstance(data, list) else 0
            except Exception:
                pass

        nb_historique = 0
        if historique_file.exists():
            try:
                data = json.loads(historique_file.read_text(encoding="utf-8"))
                nb_historique = len(data) if isinstance(data, list) else 0
            except Exception:
                pass

        # Mettre a jour le timestamp de synchro
        meta_file = DASHBOARD_DIR / ".sync_meta.json"
        meta = {}
        if meta_file.exists():
            try:
                meta = json.loads(meta_file.read_text(encoding="utf-8"))
            except Exception:
                pass
        meta["dernier_pull"] = datetime.now().isoformat()
        meta_file.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

        return jsonify({
            "status": "ok",
            "nb_ao": len(appels),
            "nb_concurrence": nb_concurrence,
            "nb_historique": nb_historique,
            "timestamp": meta["dernier_pull"],
        })
    except Exception as e:
        logger.error(f"Erreur sync pull: {e}")
        return jsonify({"status": "error", "erreur": str(e)}), 500


# --- Batch endpoints ---

@app.route("/api/batch/generer", methods=["POST"])
def api_batch_generer():
    """POST /api/batch/generer - Lance la generation batch en background."""
    data = request.get_json()
    if not data or "ao_ids" not in data:
        return jsonify({"error": "JSON body avec ao_ids requis"}), 400

    ao_ids = data["ao_ids"]
    if not ao_ids:
        return jsonify({"error": "Liste ao_ids vide"}), 400

    def _batch_bg():
        try:
            from generation_batch import generer_batch
            appels = charger_ao()
            generer_batch(ao_ids, appels, socketio=socketio)
        except Exception as e:
            logger.error(f"Erreur batch generation: {e}")
            socketio.emit("batch_progress", {"msg": f"Erreur fatale batch: {e}", "phase": "error"})

    thread = threading.Thread(target=_batch_bg, daemon=True)
    thread.start()
    return jsonify({"status": "started", "nb": len(ao_ids)})


@app.route("/api/batch/statut", methods=["POST"])
def api_batch_statut():
    """POST /api/batch/statut - Change le statut de plusieurs AO d'un coup."""
    data = request.get_json()
    if not data or "ao_ids" not in data or "statut" not in data:
        return jsonify({"error": "JSON body avec ao_ids et statut requis"}), 400

    ao_ids = data["ao_ids"]
    nouveau_statut = data["statut"]
    if nouveau_statut not in ("nouveau", "analyse", "candidature", "soumis", "ignore", "gagne", "perdu"):
        return jsonify({"error": f"Statut invalide: {nouveau_statut}"}), 400

    appels = charger_ao()
    modifies = 0
    for ao in appels:
        if ao.get("id") in ao_ids:
            ao["statut"] = nouveau_statut
            modifies += 1
    sauvegarder_ao(appels)

    # Recalibrer scoring si gagne/perdu
    if nouveau_statut in ("gagne", "perdu"):
        try:
            from scoring_predictif import calibrer_auto
            calibrer_auto()
        except Exception as e:
            logger.warning(f"Scoring predictif batch: {e}")

    return jsonify({"status": "ok", "modifies": modifies, "statut": nouveau_statut})


# --- Feature : Regenerer une piece individuelle ---

@app.route("/api/dossiers/<path:nom>/regenerer/<path:fichier>", methods=["POST"])
def api_regenerer_piece(nom, fichier):
    """POST /api/dossiers/<nom>/regenerer/<fichier> - Regenere une piece du dossier."""
    def _regen_bg():
        try:
            from generateur_render import regenerer_piece
            socketio.emit("veille_log", {"msg": f"Regeneration: {fichier}..."})
            result = regenerer_piece(nom, fichier)
            if result["success"]:
                socketio.emit("veille_log", {"msg": f"Regenere: {fichier} ({result['nb_mots']} mots, {result['duree_sec']}s)"})
            else:
                socketio.emit("veille_log", {"msg": f"Erreur regen: {result['erreur']}"})
            socketio.emit("regen_complete", {"dossier": nom, "fichier": fichier, "result": result})
        except Exception as e:
            socketio.emit("veille_log", {"msg": f"Erreur regen: {e}"})

    thread = threading.Thread(target=_regen_bg, daemon=True)
    thread.start()
    return jsonify({"status": "started", "fichier": fichier})


# --- Feature : Editeur inline markdown ---

@app.route("/dossiers/<path:nom>/editer/<path:fichier>", methods=["GET", "POST"])
def editer_fichier(nom, fichier):
    """Editeur inline pour les fichiers .md du dossier."""
    # Trouver le fichier
    fichier_path = None
    for base in [RESULTATS_DIR, DOSSIERS_GENERES_DIR]:
        candidate = base / nom / fichier
        if candidate.exists():
            fichier_path = candidate
            break

    if not fichier_path or not fichier_path.suffix.lower() == ".md":
        return redirect(url_for("detail_dossier", nom=nom))

    if request.method == "POST":
        contenu = request.form.get("contenu", "")

        # Sauvegarder une copie de backup avant modification
        import shutil
        backup_path = fichier_path.with_suffix(".md.bak")
        if fichier_path.exists():
            shutil.copy2(str(fichier_path), str(backup_path))

        fichier_path.write_text(contenu, encoding="utf-8")

        # Re-convertir en DOCX si possible
        try:
            from export_docx_render import markdown_to_docx, DOCX_DISPONIBLE
            if DOCX_DISPONIBLE:
                docx_path = fichier_path.with_suffix(".docx")
                fiche_ao = fichier_path.parent / "fiche_ao.json"
                titre_ao = ""
                if fiche_ao.exists():
                    titre_ao = json.loads(fiche_ao.read_text(encoding="utf-8")).get("titre", "")
                titre_doc = fichier.replace(".md", "").split("_", 1)[-1].replace("_", " ").title()
                markdown_to_docx(contenu, str(docx_path), titre_document=titre_ao, sous_titre=titre_doc)
        except Exception as e:
            logger.warning(f"Re-conversion DOCX apres edit: {e}")

        return redirect(url_for("servir_fichier", nom=nom, fichier=fichier))

    contenu = fichier_path.read_text(encoding="utf-8")
    contenu_html = _convertir_md_en_html(contenu)
    return render_template("editeur.html", nom_dossier=nom, nom_fichier=fichier,
                           contenu=contenu, contenu_html=contenu_html)


# --- Feature : Export PDF natif ---

@app.route("/dossiers/<path:nom>/pdf")
def export_dossier_pdf(nom):
    """Exporte le dossier complet en PDF via xhtml2pdf."""
    from flask import send_file as _send_file

    # Collecter tous les fichiers .md du dossier
    dossier_path = None
    for base in [RESULTATS_DIR, DOSSIERS_GENERES_DIR]:
        candidate = base / nom
        if candidate.exists():
            dossier_path = candidate
            break

    if not dossier_path:
        return "Dossier introuvable", 404

    # Construire le HTML complet
    html_parts = [f"""<!DOCTYPE html>
<html><head><meta charset="utf-8">
<title>Dossier {nom}</title>
<style>
    body {{ font-family: 'Helvetica', 'Arial', sans-serif; font-size: 11pt; line-height: 1.6; color: #1e293b; margin: 2cm; }}
    h1 {{ color: #1e40af; font-size: 20pt; border-bottom: 3px solid #1e40af; padding-bottom: 8px; page-break-before: always; }}
    h1:first-of-type {{ page-break-before: avoid; }}
    h2 {{ color: #1e40af; font-size: 15pt; border-bottom: 1px solid #cbd5e1; padding-bottom: 4px; }}
    h3 {{ color: #334155; font-size: 13pt; }}
    table {{ width: 100%; border-collapse: collapse; margin: 10px 0; }}
    th {{ background: #1e40af; color: white; padding: 8px; text-align: left; font-size: 10pt; }}
    td {{ padding: 6px 8px; border-bottom: 1px solid #e2e8f0; font-size: 10pt; }}
    tr:nth-child(even) td {{ background: #f8fafc; }}
    .cover {{ text-align: center; padding: 100px 40px; }}
    .cover h1 {{ font-size: 28pt; border: none; page-break-before: avoid; }}
    .cover p {{ font-size: 14pt; color: #64748b; }}
    .almera-brand {{ color: #1e40af; font-size: 16pt; font-weight: bold; }}
    @page {{ margin: 2cm; @bottom-center {{ content: "Page " counter(page) " / " counter(pages); font-size: 9pt; color: #94a3b8; }} }}
</style></head><body>
<div class="cover">
    <p class="almera-brand">ALMERA - AI MENTOR</p>
    <h1>{nom.replace('_', ' ')}</h1>
    <p>Dossier de candidature</p>
    <p style="font-size:11pt; margin-top:40px;">25 rue Campagne Premiere, 75014 Paris | contact@almera.one | almera.one</p>
</div>"""]

    md_files = sorted([f for f in dossier_path.iterdir() if f.suffix == ".md"])
    for md_file in md_files:
        contenu = md_file.read_text(encoding="utf-8")
        html = _convertir_md_en_html(contenu)
        html_parts.append(f'<div class="document-section">{html}</div>')

    html_parts.append("</body></html>")
    html_complet = "\n".join(html_parts)

    # Generer le PDF
    try:
        from xhtml2pdf import pisa
        pdf_path = dossier_path / f"{nom}.pdf"
        with open(str(pdf_path), "wb") as f:
            pisa_status = pisa.CreatePDF(html_complet, dest=f)
        if pisa_status.err:
            return "Erreur generation PDF", 500
        return _send_file(str(pdf_path), as_attachment=True, download_name=f"{nom}.pdf")
    except ImportError:
        # Fallback : servir le HTML pour impression navigateur
        return html_complet, 200, {"Content-Type": "text/html; charset=utf-8"}


# --- Feature : Historique des generations ---

@app.route("/api/generations")
def api_generations():
    """GET /api/generations - Historique des generations."""
    try:
        from generateur_render import charger_historique_generations
        log = charger_historique_generations()
        # Stats agregees
        cout_total = sum(e.get("cout_estime_usd", 0) for e in log)
        mots_total = sum(e.get("nb_mots", 0) for e in log)
        return jsonify({
            "historique": log[-50:],  # 50 derniers
            "stats": {
                "nb_generations": len(log),
                "cout_total_usd": round(cout_total, 2),
                "mots_total": mots_total,
            }
        })
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# --- Feature : Score qualite dans detail dossier ---

@app.route("/api/dossiers/<path:nom>/qualite")
def api_qualite_dossier(nom):
    """GET /api/dossiers/<nom>/qualite - Score qualite du dossier."""
    for base in [RESULTATS_DIR, DOSSIERS_GENERES_DIR]:
        dossier_path = base / nom
        if not dossier_path.exists():
            continue

        result = {}
        # Review auto
        review_path = dossier_path / "review_auto.json"
        if review_path.exists():
            try:
                result["review"] = json.loads(review_path.read_text(encoding="utf-8"))
            except Exception:
                pass

        # Conformite RC
        rc_path = dossier_path / "conformite_rc.json"
        if rc_path.exists():
            try:
                result["conformite_rc"] = json.loads(rc_path.read_text(encoding="utf-8"))
            except Exception:
                pass

        # Coherence
        coh_path = dossier_path / "coherence_check.json"
        if coh_path.exists():
            try:
                result["coherence"] = json.loads(coh_path.read_text(encoding="utf-8"))
            except Exception:
                pass

        if result:
            # Score global = moyenne ponderee
            scores = []
            if result.get("review", {}).get("score_qualite"):
                scores.append(("Review IA", result["review"]["score_qualite"], 3))
            if result.get("conformite_rc", {}).get("score"):
                scores.append(("Conformite RC", result["conformite_rc"]["score"], 2))
            if result.get("coherence", {}).get("score"):
                scores.append(("Coherence", result["coherence"]["score"], 1))

            if scores:
                total_poids = sum(s[2] for s in scores)
                score_global = round(sum(s[1] * s[2] for s in scores) / total_poids)
                result["score_global"] = score_global
                result["details_scores"] = [{"nom": s[0], "score": s[1]} for s in scores]

            return jsonify(result)

    return jsonify({"error": "Dossier introuvable"}), 404


# --- Feature : Diff visuel avant/apres ---

@app.route("/dossiers/<path:nom>/diff/<path:fichier>")
def diff_fichier(nom, fichier):
    """Affiche le diff entre la version actuelle et le backup (.bak)."""
    fichier_path = None
    for base in [RESULTATS_DIR, DOSSIERS_GENERES_DIR]:
        candidate = base / nom / fichier
        if candidate.exists():
            fichier_path = candidate
            break

    if not fichier_path:
        return redirect(url_for("detail_dossier", nom=nom))

    backup_path = fichier_path.with_suffix(".md.bak")
    if not backup_path.exists():
        return render_template("diff_view.html", nom_dossier=nom, nom_fichier=fichier,
                               ancien="", nouveau=fichier_path.read_text(encoding="utf-8"),
                               has_diff=False)

    ancien = backup_path.read_text(encoding="utf-8")
    nouveau = fichier_path.read_text(encoding="utf-8")

    # Calculer les lignes de diff
    import difflib
    diff_lines = list(difflib.unified_diff(
        ancien.splitlines(keepends=True),
        nouveau.splitlines(keepends=True),
        fromfile="Avant",
        tofile="Apres",
        lineterm=""
    ))

    return render_template("diff_view.html", nom_dossier=nom, nom_fichier=fichier,
                           ancien=ancien, nouveau=nouveau,
                           diff_lines=diff_lines, has_diff=True)


@app.route("/dossiers/<path:nom>/diff/<path:fichier>/restaurer", methods=["POST"])
def restaurer_fichier(nom, fichier):
    """Restaure le fichier depuis le backup."""
    for base in [RESULTATS_DIR, DOSSIERS_GENERES_DIR]:
        fichier_path = base / nom / fichier
        backup_path = fichier_path.with_suffix(".md.bak")
        if backup_path.exists():
            import shutil
            shutil.copy2(str(backup_path), str(fichier_path))
            backup_path.unlink()
            break
    return redirect(url_for("detail_dossier", nom=nom))


# --- Feature : Mode revision guidee ---

REVISION_FILE = DASHBOARD_DIR / "revisions_guidees.json"


def _charger_revisions() -> dict:
    if REVISION_FILE.exists():
        try:
            return json.loads(REVISION_FILE.read_text(encoding="utf-8"))
        except (json.JSONDecodeError, OSError):
            pass
    return {}


def _sauvegarder_revisions(data: dict):
    REVISION_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


@app.route("/dossiers/<path:nom>/revision")
def revision_guidee(nom):
    """Page de revision guidee du dossier."""
    dossiers = lister_dossiers()
    dossier = next((d for d in dossiers if d["nom"] == nom), None)
    if not dossier:
        return redirect(url_for("liste_dossiers"))

    # Charger l'etat de revision
    revisions = _charger_revisions()
    etat = revisions.get(nom, {})

    # Construire la checklist a partir des fichiers
    checklist = []
    for f in sorted(dossier["fichiers"]):
        if not f.endswith(".md"):
            continue
        base_name = f.replace(".md", "").split("_", 1)[-1].replace("_", " ").title()
        checklist.append({
            "fichier": f,
            "label": base_name,
            "valide": etat.get(f, {}).get("valide", False),
            "commentaire": etat.get(f, {}).get("commentaire", ""),
            "date_revision": etat.get(f, {}).get("date_revision", ""),
        })

    # Charger la qualite si dispo
    qualite = {}
    for base in [RESULTATS_DIR, DOSSIERS_GENERES_DIR]:
        review_path = base / nom / "review_auto.json"
        if review_path.exists():
            try:
                qualite = json.loads(review_path.read_text(encoding="utf-8"))
            except Exception:
                pass
            break

    return render_template("revision_guidee.html", dossier=dossier, checklist=checklist,
                           qualite=qualite, etat=etat)


@app.route("/dossiers/<path:nom>/revision", methods=["POST"])
def revision_guidee_post(nom):
    """Enregistre l'etat de la revision guidee."""
    revisions = _charger_revisions()
    if nom not in revisions:
        revisions[nom] = {}

    fichier = request.form.get("fichier", "")
    action = request.form.get("action", "")

    if action == "valider":
        revisions[nom][fichier] = {
            "valide": True,
            "commentaire": request.form.get("commentaire", ""),
            "date_revision": datetime.now().strftime("%Y-%m-%d %H:%M"),
        }
    elif action == "rejeter":
        revisions[nom][fichier] = {
            "valide": False,
            "commentaire": request.form.get("commentaire", ""),
            "date_revision": datetime.now().strftime("%Y-%m-%d %H:%M"),
        }
    elif action == "valider_tout":
        dossiers = lister_dossiers()
        dossier = next((d for d in dossiers if d["nom"] == nom), None)
        if dossier:
            for f in dossier["fichiers"]:
                if f.endswith(".md"):
                    revisions[nom][f] = {
                        "valide": True,
                        "commentaire": "Valide en bloc",
                        "date_revision": datetime.now().strftime("%Y-%m-%d %H:%M"),
                    }

    _sauvegarder_revisions(revisions)
    return redirect(url_for("revision_guidee", nom=nom))


# --- Feature : CRM acheteur enrichi (stats) ---

@app.route("/api/crm/<path:nom>/stats")
def api_crm_stats(nom):
    """GET /api/crm/<nom>/stats - Statistiques enrichies d'un acheteur."""
    from crm_acheteurs import get_fiche
    fiche = get_fiche(nom, is_cle=True)
    if not fiche:
        return jsonify({"error": "Acheteur introuvable"}), 404

    historique = fiche.get("historique_ao", [])
    nb_ao = len(historique)
    nb_gagnes = sum(1 for h in historique if h.get("statut") == "gagne")
    nb_perdus = sum(1 for h in historique if h.get("statut") == "perdu")
    nb_soumis = sum(1 for h in historique if h.get("statut") in ("soumis", "gagne", "perdu"))
    montant_total = sum(h.get("montant") or 0 for h in historique if h.get("montant"))
    montant_moyen = round(montant_total / max(1, nb_ao))
    win_rate = round(nb_gagnes / max(1, nb_soumis) * 100) if nb_soumis else 0

    # Concurrents frequents (depuis concurrence.json)
    concurrents_freq = []
    try:
        conc_file = DASHBOARD_DIR / "concurrence.json"
        if conc_file.exists():
            attributions = json.loads(conc_file.read_text(encoding="utf-8"))
            acheteur_nom = fiche.get("nom", "").lower()
            concurrents_count = {}
            for a in attributions:
                if acheteur_nom in (a.get("acheteur") or "").lower():
                    t = a.get("titulaire", "").strip()
                    if t and "almera" not in t.lower() and "ai mentor" not in t.lower():
                        concurrents_count[t] = concurrents_count.get(t, 0) + 1
            concurrents_freq = sorted(concurrents_count.items(), key=lambda x: -x[1])[:10]
    except Exception:
        pass

    # Mots-cles frequents dans les titres
    mots_freq = {}
    stop_words = {"de", "des", "du", "la", "le", "les", "et", "en", "pour", "un", "une", "d", "l", "a"}
    for h in historique:
        for mot in (h.get("titre") or "").lower().split():
            mot = mot.strip(".,;:()[]'-/")
            if len(mot) > 3 and mot not in stop_words:
                mots_freq[mot] = mots_freq.get(mot, 0) + 1
    top_mots = sorted(mots_freq.items(), key=lambda x: -x[1])[:15]

    return jsonify({
        "nb_ao": nb_ao,
        "nb_gagnes": nb_gagnes,
        "nb_perdus": nb_perdus,
        "nb_soumis": nb_soumis,
        "win_rate": win_rate,
        "montant_total": montant_total,
        "montant_moyen": montant_moyen,
        "concurrents_frequents": [{"nom": n, "count": c} for n, c in concurrents_freq],
        "mots_cles": [{"mot": m, "count": c} for m, c in top_mots],
    })


# --- Feature : Import DCE drag & drop ---

@app.route("/api/ao/<path:ao_id>/upload-dce", methods=["POST"])
def api_upload_dce(ao_id):
    """POST /api/ao/<id>/upload-dce - Upload de fichiers DCE par drag & drop."""
    fichiers = request.files.getlist("fichiers")
    if not fichiers:
        return jsonify({"error": "Aucun fichier"}), 400

    # Creer le dossier DCE
    clean_id = ao_id.replace("/", "_").replace("\\", "_")
    dossier_dce = DOSSIERS_GENERES_DIR / f"DCE_{clean_id}"
    dossier_dce.mkdir(parents=True, exist_ok=True)

    fichiers_sauves = []
    for f in fichiers:
        if f.filename:
            import re as _re
            safe_name = _re.sub(r"[^a-zA-Z0-9_.-]", "_", f.filename)
            f.save(str(dossier_dce / safe_name))
            fichiers_sauves.append(safe_name)

    # Lancer l'analyse DCE en arriere-plan
    def _analyse_bg():
        try:
            from analyse_semantique_dce import analyser_dce_complet_ia
            appels = charger_ao()
            ao_dict = next((a for a in appels if a.get("id") == ao_id), None)
            if ao_dict:
                result = analyser_dce_complet_ia(dossier_dce, ao_dict)
                socketio.emit("veille_log", {"msg": f"DCE analyse: score adequation {result.get('score_adequation', '?')}/100"})
                socketio.emit("dce_analyse_complete", {"ao_id": ao_id, "result": result})
        except Exception as e:
            logger.warning(f"Analyse DCE auto upload: {e}")
            socketio.emit("veille_log", {"msg": f"Analyse DCE: {e}"})

    thread = threading.Thread(target=_analyse_bg, daemon=True)
    thread.start()

    return jsonify({
        "success": True,
        "fichiers": fichiers_sauves,
        "dossier": f"DCE_{clean_id}",
        "analyse_lancee": True,
    })


# ===================================================================
# Feature 5+13 : Authentification multi-utilisateur (Flask-Login)
# ===================================================================

try:
    from auth import init_auth, check_login, login_required as auth_required
    from flask_login import current_user, login_user, logout_user
    init_auth(app)
    AUTH_ENABLED = True
    logger.info("Authentification activee")
except ImportError:
    AUTH_ENABLED = False
    logger.info("Auth non disponible (flask-login manquant) - mode ouvert")

    # Dummy decorateur
    def auth_required(f):
        return f
    class _FakeUser:
        is_authenticated = True
        nom = "Admin"
        role = "admin"
        is_admin = True
        id = "admin"
    current_user = _FakeUser()


@app.route("/login", methods=["GET", "POST"])
def login():
    if not AUTH_ENABLED:
        return redirect(url_for("index"))
    if request.method == "POST":
        user = check_login(request.form.get("username", ""), request.form.get("password", ""))
        if user:
            login_user(user, remember=True)
            return redirect(request.args.get("next") or url_for("index"))
        from flask import flash
        flash("Identifiant ou mot de passe incorrect.", "error")
    return render_template("login.html")


@app.route("/logout")
def logout():
    if AUTH_ENABLED:
        logout_user()
    return redirect(url_for("login"))


@app.context_processor
def inject_auth():
    if AUTH_ENABLED:
        from flask_login import current_user as cu
        return {"current_user": cu, "auth_enabled": True}
    return {"current_user": current_user, "auth_enabled": False}


# ===================================================================
# Feature 3 : Self-ping anti cold-start (Render)
# ===================================================================

def _self_ping():
    """Ping l'app elle-meme pour eviter le cold start Render."""
    try:
        import httpx
        url = os.environ.get("RENDER_EXTERNAL_URL", "https://ao-hunter-dashboard.onrender.com")
        with httpx.Client(timeout=10) as client:
            resp = client.get(f"{url}/api/ping")
            logger.debug(f"Self-ping: {resp.status_code}")
    except Exception:
        pass


@app.route("/api/ping")
def api_ping():
    return jsonify({"status": "ok", "time": datetime.now().isoformat()})


# ===================================================================
# Feature 6 : Pipeline auto idempotent
# ===================================================================

PIPELINE_PROCESSED_FILE = DASHBOARD_DIR / "pipeline_processed.json"

def _charger_pipeline_processed():
    if PIPELINE_PROCESSED_FILE.exists():
        try:
            return json.loads(PIPELINE_PROCESSED_FILE.read_text(encoding="utf-8"))
        except Exception:
            pass
    return {"processed_ids": [], "last_run": None}


def _marquer_pipeline_processed(ao_id):
    data = _charger_pipeline_processed()
    if ao_id not in data["processed_ids"]:
        data["processed_ids"].append(ao_id)
    data["last_run"] = datetime.now().isoformat()
    # Garder les 500 derniers
    data["processed_ids"] = data["processed_ids"][-500:]
    PIPELINE_PROCESSED_FILE.write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")


@app.route("/api/pipeline/processed")
def api_pipeline_processed():
    return jsonify(_charger_pipeline_processed())


# ===================================================================
# Feature 11 : Webhook Slack/Teams alertes
# ===================================================================

WEBHOOK_URL = os.environ.get("AO_HUNTER_WEBHOOK_URL", "")  # Slack ou Teams


def envoyer_webhook(message: str, urgence: str = "info"):
    """Envoie une notification via webhook Slack/Teams."""
    if not WEBHOOK_URL:
        return
    try:
        import httpx
        # Format Slack
        payload = {
            "text": f"{'🔴' if urgence == 'urgent' else '🔵'} *AO Hunter* - {message}"
        }
        with httpx.Client(timeout=10) as client:
            client.post(WEBHOOK_URL, json=payload)
    except Exception as e:
        logger.warning(f"Webhook erreur: {e}")


# ===================================================================
# Feature 12 : Scoring IA (route API)
# ===================================================================

@app.route("/api/ao/<ao_id>/score-ia", methods=["POST"])
def api_score_ia(ao_id):
    """Score un AO via Claude IA."""
    appels = charger_ao()
    ao = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao:
        return jsonify({"error": "AO non trouve"}), 404
    try:
        from scoring_ia import scorer_ao_ia
        result = scorer_ao_ia(ao)
        # Sauvegarder le score IA dans l'AO
        ao["score_ia"] = result.get("score", 0)
        ao["score_ia_justification"] = result.get("justification", "")
        ao["score_ia_recommandation"] = result.get("recommandation", "")
        sauvegarder_ao(appels)
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ===================================================================
# Feature 1 : Backup donnees vers GitHub (route + scheduler)
# ===================================================================

@app.route("/api/backup", methods=["POST"])
def api_backup():
    """Lance un backup des donnees vers GitHub."""
    try:
        from db_backup import backup_to_git
        result = backup_to_git()
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ===================================================================
# Feature 4 : Download DCE HTTP direct
# ===================================================================

@app.route("/api/ao/<ao_id>/download-dce-http", methods=["POST"])
def api_download_dce_http(ao_id):
    """Telecharge le DCE via URL directe HTTP (sans Playwright)."""
    appels = charger_ao()
    ao = next((a for a in appels if a.get("id") == ao_id), None)
    if not ao:
        return jsonify({"error": "AO non trouve"}), 404

    url = ao.get("url_profil_acheteur") or ao.get("url", "")
    if not url:
        return jsonify({"error": "Aucune URL disponible pour cet AO"}), 400

    try:
        from dce_auto import telecharger_dce_http
        result = telecharger_dce_http(ao_id, url)
        return jsonify(result)
    except Exception as e:
        return jsonify({"error": str(e)}), 500


# ===================================================================
# Feature 14 : Dashboard concurrent enrichi
# ===================================================================

@app.route("/concurrence")
def page_concurrence():
    """Dashboard concurrent avec graphiques."""
    conc_file = DASHBOARD_DIR / "concurrence.json"
    concurrents = []
    if conc_file.exists():
        try:
            concurrents = json.loads(conc_file.read_text(encoding="utf-8"))
        except Exception:
            pass

    # Agreger par concurrent
    from collections import Counter, defaultdict
    par_concurrent = defaultdict(lambda: {"nb": 0, "secteurs": Counter(), "montants": [], "mois": Counter()})
    for c in concurrents:
        nom = c.get("titulaire", "Inconnu")
        par_concurrent[nom]["nb"] += 1
        sect = c.get("secteur") or c.get("type_marche") or "Autre"
        par_concurrent[nom]["secteurs"][sect] += 1
        if c.get("montant"):
            try:
                par_concurrent[nom]["montants"].append(float(c["montant"]))
            except (ValueError, TypeError):
                pass
        date_str = c.get("date_attribution") or c.get("date_publication") or ""
        if len(date_str) >= 7:
            par_concurrent[nom]["mois"][date_str[:7]] += 1

    # Top 20
    top = sorted(par_concurrent.items(), key=lambda x: x[1]["nb"], reverse=True)[:20]
    top_data = []
    for nom, data in top:
        top_data.append({
            "nom": nom,
            "nb_marches": data["nb"],
            "montant_moyen": round(sum(data["montants"]) / len(data["montants"])) if data["montants"] else 0,
            "montant_total": round(sum(data["montants"])) if data["montants"] else 0,
            "top_secteurs": data["secteurs"].most_common(3),
            "tendance": dict(sorted(data["mois"].items())[-12:]),
        })

    # Stats globales
    total_marches = len(concurrents)
    nb_concurrents_uniques = len(par_concurrent)
    montant_total = sum(float(c.get("montant", 0)) for c in concurrents if c.get("montant"))

    return render_template("concurrence_dashboard.html",
                           top_concurrents=top_data,
                           total_marches=total_marches,
                           nb_concurrents=nb_concurrents_uniques,
                           montant_total=montant_total,
                           concurrents_json=json.dumps(top_data, ensure_ascii=False))


# ===================================================================
# Nouvelles routes : Resultats, Recherche, Resume, Validation
# ===================================================================

@app.route("/resultats")
def page_resultats():
    """Page de suivi des resultats Win/Loss."""
    try:
        from win_loss_tracker import analyser_performances
        stats = analyser_performances(periode_mois=12)
    except Exception as e:
        logger.warning(f"Win/Loss tracker: {e}")
        stats = {
            'global': {'nb_gagnes': 0, 'nb_perdus': 0, 'win_rate': 0, 'ca_gagne': 0, 'ca_moyen': 0, 'nb_abandons': 0, 'nb_sans_suite': 0, 'nb_total': 0},
            'par_type_acheteur': {}, 'par_tranche_budget': {}, 'par_source': {},
            'facteurs_succes': [], 'facteurs_echec': [], 'concurrents_frequents': [],
            'tendance': {'mois': [], 'win_rates': []}, 'recommandations': [],
        }
    # AO soumis (pour le formulaire d'enregistrement)
    aos = charger_ao()
    aos_soumis = [ao for ao in aos if ao.get('statut') in ('soumis', 'candidature', 'gagne', 'perdu')]
    return render_template("resultats_tracker.html", stats=stats, aos_soumis=aos_soumis)


@app.route("/api/resultats", methods=["POST"])
def api_enregistrer_resultat():
    """Enregistre un resultat Win/Loss."""
    try:
        from win_loss_tracker import enregistrer_resultat
        data = request.get_json() or request.form.to_dict()
        ao_id = data.get('ao_id')
        resultat = data.get('resultat')
        if not ao_id or not resultat:
            return jsonify({"error": "ao_id et resultat requis"}), 400
        details = {
            'raison_principale': data.get('raison_principale', ''),
            'montant_final': float(data['montant_final']) if data.get('montant_final') else 0,
            'concurrent_gagnant': data.get('concurrent_gagnant', ''),
        }
        entry = enregistrer_resultat(ao_id, resultat, details)
        return jsonify({"ok": True, "entry": entry})
    except Exception as e:
        logger.error(f"Erreur enregistrement resultat: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/api/resultats/stats")
def api_resultats_stats():
    """Stats Win/Loss en JSON."""
    try:
        from win_loss_tracker import analyser_performances
        return jsonify(analyser_performances(periode_mois=12))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/recherche")
def page_recherche():
    """Page de recherche globale."""
    query = request.args.get('q', '').strip()
    resultats = {'nb_total': 0, 'resultats': {'ao': [], 'dossiers': [], 'notes': [], 'crm': []}}
    suggestions = []

    if query:
        try:
            from recherche_globale import rechercher
            resultats = rechercher(query)
        except Exception as e:
            logger.warning(f"Recherche: {e}")
    else:
        try:
            from recherche_globale import suggestions_recherche
            suggestions = suggestions_recherche()
        except Exception:
            suggestions = ['formation IA', 'intelligence artificielle', 'Qualiopi']

    return render_template("recherche.html", query=query, resultats=resultats, suggestions=suggestions)


@app.route("/api/recherche")
def api_recherche():
    """API de recherche globale."""
    query = request.args.get('q', '').strip()
    if not query:
        return jsonify({"error": "Parametre q requis"}), 400
    try:
        from recherche_globale import rechercher
        return jsonify(rechercher(query))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/resume")
def page_resume():
    """Page du resume quotidien."""
    try:
        from resume_quotidien import generer_resume
        resume = generer_resume()
    except Exception as e:
        logger.warning(f"Resume quotidien: {e}")
        resume = {
            'date': datetime.now().strftime('%A %d %B %Y'),
            'resume_texte': 'Resume indisponible',
            'kpis': {'total_ao': 0, 'score_moyen': 0, 'nb_dossiers': 0, 'nb_urgents': 0, 'win_rate': 0, 'nb_gagnes': 0, 'nb_perdus': 0},
            'nouveaux_ao': [], 'deadlines_imminentes': [], 'actions_prioritaires': [],
            'opportunites': [], 'pipeline': {},
        }
    return render_template("resume_quotidien.html", resume=resume)


@app.route("/api/resume")
def api_resume():
    """API resume quotidien en JSON."""
    try:
        from resume_quotidien import generer_resume
        return jsonify(generer_resume())
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/ao/<ao_id>/validation")
def api_validation_dossier(ao_id):
    """Validation de completude d'un dossier genere."""
    try:
        from generateur_render import valider_dossier_complet
        ao = next((a for a in charger_ao() if a.get('id') == ao_id), None)
        if not ao:
            return jsonify({"error": "AO non trouve"}), 404
        # Trouver le dossier
        dossier_name = f"AO_{ao_id.replace('/', '_').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d')}"
        dossier_path = DOSSIERS_GENERES_DIR / dossier_name
        # Chercher le dossier le plus recent pour cet AO
        if not dossier_path.exists():
            prefix = f"AO_{ao_id.replace('/', '_').replace(' ', '_')}"
            candidates = sorted(DOSSIERS_GENERES_DIR.glob(f"{prefix}*"), reverse=True)
            if candidates:
                dossier_path = candidates[0]
            else:
                return jsonify({"error": "Aucun dossier genere pour cet AO"}), 404
        result = valider_dossier_complet(ao, dossier_path)
        return jsonify(result)
    except Exception as e:
        logger.error(f"Validation dossier: {e}")
        return jsonify({"error": str(e)}), 500


@app.route("/carte")
def page_carte_regions():
    """Carte des opportunites par region."""
    aos = charger_ao()
    from collections import defaultdict
    regions_data = defaultdict(lambda: {'ao_list': [], 'scores': []})

    for ao in aos:
        region = ao.get('region', '') or ao.get('lieu_execution', '') or ''
        if not region or region.lower() in ('non precise', 'non precis', ''):
            continue
        score = round((ao.get('score_pertinence', 0) or 0) * 100)
        regions_data[region]['ao_list'].append({'id': ao.get('id', ''), 'titre': ao.get('titre', ''), 'score': score})
        regions_data[region]['scores'].append(score)

    regions = []
    for nom, data in regions_data.items():
        scores = data['scores']
        regions.append({
            'nom': nom,
            'nb_ao': len(data['ao_list']),
            'score_moyen': round(sum(scores) / len(scores)) if scores else 0,
            'top_ao': sorted(data['ao_list'], key=lambda x: x['score'], reverse=True)[:5],
        })
    regions.sort(key=lambda x: x['nb_ao'], reverse=True)

    top_region = regions[0] if regions else None
    total_ao = sum(r['nb_ao'] for r in regions)

    return render_template("carte_regions.html", regions=regions, top_region=top_region, total_ao=total_ao)


@app.route("/api/ao/<ao_id>/score-predictif")
def api_score_predictif(ao_id):
    """Score predictif de gain base sur l'historique."""
    try:
        from win_loss_tracker import score_predictif_ao
        ao = next((a for a in charger_ao() if a.get('id') == ao_id), None)
        if not ao:
            return jsonify({"error": "AO non trouve"}), 404
        return jsonify(score_predictif_ao(ao))
    except Exception as e:
        return jsonify({"error": str(e)}), 500


@app.route("/api/batch-generer", methods=["POST"])
def api_batch_generer_smart():
    """Generation batch intelligente avec priorite automatique."""
    data = request.get_json() or {}
    max_dossiers = min(int(data.get('max', 3)), 5)  # Max 5 par batch
    score_min = float(data.get('score_min', 0.6))

    aos = charger_ao()

    # Filtrer AO eligibles (pas deja generes, score suffisant)
    dossiers_index = {}
    dossiers_index_path = DASHBOARD_DIR / "dossiers_index.json"
    if dossiers_index_path.exists():
        try:
            idx = json.loads(dossiers_index_path.read_text(encoding="utf-8"))
            if isinstance(idx, list):
                dossiers_index = {d.get('ao_id', ''): d for d in idx}
            elif isinstance(idx, dict):
                dossiers_index = idx
        except Exception:
            pass

    eligibles = []
    for ao in aos:
        if ao.get('statut', 'nouveau') in ('gagne', 'perdu', 'ignore', 'abandon'):
            continue
        if ao.get('id', '') in dossiers_index:
            continue
        score = ao.get('score_pertinence', 0) or 0
        if score < score_min:
            continue

        # Score de priorite: pertinence × urgence
        priorite = score
        dl = ao.get('date_limite', '')
        if dl:
            try:
                dt = datetime.fromisoformat(str(dl).replace('Z', '+00:00')).replace(tzinfo=None)
                jours = (dt - datetime.now()).days
                if 0 < jours <= 7:
                    priorite *= 1.5
                elif 7 < jours <= 14:
                    priorite *= 1.2
                elif jours <= 0:
                    continue  # Deadline passee
            except Exception:
                pass

        eligibles.append({'ao': ao, 'priorite': priorite})

    eligibles.sort(key=lambda x: x['priorite'], reverse=True)
    selection = eligibles[:max_dossiers]

    if not selection:
        return jsonify({"ok": True, "message": "Aucun AO eligible", "generes": 0, "details": []})

    # Lancer les generations en arriere-plan
    def _batch_worker():
        resultats = []
        for i, item in enumerate(selection):
            ao = item['ao']
            try:
                socketio.emit("batch_progress", {
                    "phase": "generation",
                    "message": f"Generation {i+1}/{len(selection)}: {ao.get('titre', '')[:50]}",
                    "nb_complete": i,
                    "nb_total": len(selection),
                })
                from generateur_render import generer_dossier_complet
                result = generer_dossier_complet(ao)
                resultats.append({"ao_id": ao.get('id'), "ok": True, "dossier": result.get('nom', '')})
            except Exception as e:
                resultats.append({"ao_id": ao.get('id'), "ok": False, "erreur": str(e)})

        socketio.emit("batch_progress", {
            "phase": "termine",
            "message": f"Batch termine: {sum(1 for r in resultats if r['ok'])}/{len(resultats)} dossiers generes",
            "nb_complete": len(selection),
            "nb_total": len(selection),
        })

    thread = threading.Thread(target=_batch_worker, daemon=True)
    thread.start()

    return jsonify({
        "ok": True,
        "message": f"Generation batch lancee pour {len(selection)} AO",
        "generes": len(selection),
        "details": [{"ao_id": s['ao'].get('id'), "titre": s['ao'].get('titre', '')[:60], "priorite": round(s['priorite'] * 100)} for s in selection],
    })


# ===================================================================
# Scheduler : ajouter self-ping + backup periodique
# ===================================================================

def _resume_quotidien_auto():
    """Genere et envoie le resume quotidien par email (brouillon Gmail)."""
    try:
        from resume_quotidien import generer_email_resume
        email = generer_email_resume()
        logger.info(f"Resume quotidien genere: {email['sujet']}")

        # Creer brouillon Gmail si SMTP dispo
        smtp_password = os.environ.get("SMTP_PASSWORD", "")
        if smtp_password:
            try:
                import imaplib
                from email.mime.text import MIMEText
                msg = MIMEText(email['corps'], 'plain', 'utf-8')
                msg['Subject'] = email['sujet']
                msg['From'] = 'mickael.bertolla@gmail.com'
                msg['To'] = email['destinataire']

                imap = imaplib.IMAP4_SSL("imap.gmail.com", 993)
                imap.login('mickael.bertolla@gmail.com', smtp_password)
                imap.append('[Gmail]/Brouillons', '', None, msg.as_bytes())
                imap.logout()
                logger.info("Resume quotidien: brouillon Gmail cree")
            except Exception as e:
                logger.warning(f"Resume quotidien: brouillon non cree: {e}")
        else:
            logger.info("Resume quotidien: pas de SMTP_PASSWORD, brouillon non cree")
    except Exception as e:
        logger.error(f"Erreur resume quotidien: {e}")


def init_scheduler_enhanced():
    """Version amelioree du scheduler avec self-ping + backup."""
    global _scheduler
    try:
        from apscheduler.schedulers.background import BackgroundScheduler
        from apscheduler.triggers.interval import IntervalTrigger
        from apscheduler.triggers.cron import CronTrigger

        if _scheduler is None:
            _scheduler = BackgroundScheduler(daemon=True)

        # Self-ping toutes les 13 minutes (anti cold-start)
        _scheduler.add_job(
            func=_self_ping,
            trigger=IntervalTrigger(minutes=13),
            id="self_ping",
            name="Self-ping anti cold-start",
            replace_existing=True,
        )

        # Backup donnees toutes les 6h
        _scheduler.add_job(
            func=lambda: __import__("db_backup").backup_to_git() if os.environ.get("RENDER") else None,
            trigger=IntervalTrigger(hours=6),
            id="auto_backup",
            name="Backup donnees GitHub",
            replace_existing=True,
        )

        # Resume quotidien chaque jour a 8h30
        _scheduler.add_job(
            func=_resume_quotidien_auto,
            trigger=CronTrigger(hour=8, minute=30),
            id="resume_quotidien",
            name="Resume quotidien 8h30",
            replace_existing=True,
        )

        if not _scheduler.running:
            _scheduler.start()
        logger.info("Scheduler enhanced: self-ping + backup ajoutes")
    except Exception as e:
        logger.warning(f"Scheduler enhanced non disponible: {e}")


# Demarrer le scheduler sur Render (pas en local, la tache planifiee Windows s'en charge)
if os.environ.get("RENDER"):
    init_scheduler()
    init_scheduler_enhanced()
else:
    # En local, juste le self-ping pour le dev
    try:
        init_scheduler_enhanced()
    except Exception:
        pass


if __name__ == "__main__":
    port = int(os.environ.get("PORT", 5000))
    debug = os.environ.get("RENDER") is None  # debug only in local
    print(f"AO Hunter Dashboard - http://localhost:{port}")
    socketio.run(app, debug=debug, use_reloader=False, host="0.0.0.0", port=port, allow_unsafe_werkzeug=True)
